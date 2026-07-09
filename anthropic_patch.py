import re
import os

with open('app.py', 'r', encoding='utf-8') as f:
    content = f.read()

anthropic_code = """
# ==========================================
# Anthropic Claude Compatibility Layer
# ==========================================

def call_anthropic_api(model, messages, temperature=0, timeout=None):
    import anthropic
    config = get_config()
    api_key = config.get('ANTHROPIC_API_KEY', '')
    if not api_key:
        raise ValueError("Anthropic API key not configured. Set ANTHROPIC_API_KEY in settings.")

    client = anthropic.Anthropic(api_key=api_key, timeout=timeout or float(config.get('ANTHROPIC_TIMEOUT_SECONDS', 240)))

    system_prompt = ""
    anthropic_messages = []

    for msg in messages:
        if msg['role'] == 'system':
            system_prompt += msg['content'] + "\\n"
        elif msg['role'] in ['user', 'assistant']:
            if isinstance(msg['content'], list):
                new_content = []
                for item in msg['content']:
                    if item['type'] == 'text':
                        new_content.append({"type": "text", "text": item['text']})
                    elif item['type'] == 'image_url':
                        url = item['image_url']['url']
                        media_type, base64_data = url.split(";base64,")
                        media_type = media_type.replace("data:", "")
                        new_content.append({
                            "type": "image",
                            "source": {
                                "type": "base64",
                                "media_type": media_type,
                                "data": base64_data
                            }
                        })
                anthropic_messages.append({"role": msg['role'], "content": new_content})
            else:
                anthropic_messages.append({"role": msg['role'], "content": msg['content']})

    response = client.messages.create(
        model=model or "claude-3-5-sonnet-latest",
        system=system_prompt.strip() if system_prompt else anthropic.NOT_GIVEN,
        messages=anthropic_messages,
        max_tokens=4096,
        temperature=temperature
    )
    return response.content[0].text

def extract_clean_text_with_anthropic(image, page_number, model):
    data_url = prepare_openai_compatible_image_data_url(image)

    message = call_anthropic_api(
        model,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            "OCR this document page and clean the result. Preserve Bangla and English exactly. "
                            "Keep question numbers, options, formulas, answer markers such as 'উ. ক', and explanations. "
                            "Do not summarize. Return only cleaned page text."
                        ),
                    },
                    {"type": "image_url", "image_url": {"url": data_url}},
                ],
            }
        ],
        temperature=0,
    )

    if message:
        return message
    raise RuntimeError(f"No text returned by Anthropic engine for page {page_number}.")

def transform_page_image_to_mcq_dataset_with_anthropic(image, page_number, model, page_text=None):
    data_url = prepare_openai_compatible_image_data_url(image)

    message = call_anthropic_api(
        model,
        messages=[
            {"role": "system", "content": MCQ_DATASET_SYSTEM_PROMPT},
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            f"OCR, clean, and structure all complete MCQs visible on page {page_number}. "
                            "Use both the cleaned OCR text below and the page image directly for OCR, options, formulas, "
                            "answer markers, and explanations. "
                            "Do not omit visible numbered questions. Right-side printed markers like 'উ. ক', 'উ. খ', "
                            "'উ: গ', 'উত্তর: ঘ' usually denote the correct answer option.\\n\\n"
                            f"=== Cleaned OCR Text (Page {page_number}) ===\\n{page_text or 'No text provided.'}"
                        ),
                    },
                    {"type": "image_url", "image_url": {"url": data_url}},
                ],
            }
        ],
        temperature=0,
    )

    return normalize_mcq_dataset(json.loads(strip_json_code_fence(message)))

def extract_missing_mcqs_with_anthropic(image, page_number, page_text, existing_dataset, model):
    data_url = prepare_openai_compatible_image_data_url(image)
    existing_mcqs_json = json.dumps(existing_dataset.get("mcqs", []), ensure_ascii=False, indent=2)

    message = call_anthropic_api(
        model,
        messages=[
            {"role": "system", "content": MCQ_DATASET_SYSTEM_PROMPT},
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            f"You previously extracted MCQs from page {page_number}, but some are missing. "
                            "Look at the page image and the extracted text again. "
                            f"These are the MCQs you already found:\\n```json\\n{existing_mcqs_json}\\n```\\n\\n"
                            "Extract ONLY the MCQs that are present in the image/text but MISSING from the JSON above. "
                            "Do not repeat already extracted MCQs. "
                            f"=== Cleaned OCR Text (Page {page_number}) ===\\n{page_text or 'No text provided.'}"
                        ),
                    },
                    {"type": "image_url", "image_url": {"url": data_url}},
                ],
            }
        ],
        temperature=0,
    )

    return normalize_mcq_dataset(json.loads(strip_json_code_fence(message)))

def process_with_anthropic(file_path, output_format, job_id=None, page_range=None):
    config = get_config()
    def update_progress(percentage, status):
        if job_id:
            job_progress[job_id] = {"status": status, "percentage": percentage}

    def check_controls():
        if job_id and job_id in job_controls:
            if job_controls[job_id].get('cancel_flag'):
                raise InterruptedError("Job cancelled by user")
            job_controls[job_id]['pause_event'].wait()

    update_progress(5, "Initializing Anthropic engine...")
    model = config.get('ANTHROPIC_MODEL', 'claude-3-5-sonnet-latest')

    total_pdf_pages = 0
    if is_pdf_file(file_path):
        doc = fitz.open(file_path)
        total_pdf_pages = len(doc)
        doc.close()

    page_indices = parse_page_range(page_range, total_pdf_pages) if is_pdf_file(file_path) else None
    page_images = get_document_page_images(file_path, dpi=int(config.get('OCR_DPI', 300)), page_indices=page_indices)
    total_pages = len(page_images)

    if output_format == "json":
        page_datasets = []
        for idx, (page_number, image) in enumerate(page_images):
            check_controls()
            update_progress(
                5 + int((idx / total_pages) * 90),
                f"Processing page {page_number} of {total_pages}..."
            )
            page_text = extract_clean_text_with_anthropic(image, page_number, model)
            page_dataset = transform_page_image_to_mcq_dataset_with_anthropic(
                image, page_number, model, page_text=page_text
            )
            expected_count = get_longest_consecutive_question_count(page_text)
            if expected_count and len(page_dataset.get("mcqs", [])) < expected_count:
                update_progress(
                    5 + int((idx / total_pages) * 90) + 2,
                    f"Repairing page {page_number} (found {len(page_dataset.get('mcqs', []))}/{expected_count})..."
                )
                missing_dataset = extract_missing_mcqs_with_anthropic(
                    image, page_number, page_text, page_dataset, model
                )
                page_dataset = merge_mcq_datasets([page_dataset, missing_dataset])
            page_datasets.append(page_dataset)

        dataset = merge_mcq_datasets(page_datasets)
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], f'anthropic_{job_id or uuid.uuid4().hex}.json')
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(dataset, f, ensure_ascii=False, indent=2)
        preview_text = json.dumps(dataset, ensure_ascii=False, indent=2)
    else:
        extracted_text = []
        for idx, (page_number, image) in enumerate(page_images):
            update_progress(
                5 + int((idx / total_pages) * 90),
                f"Processing page {page_number} of {total_pages}..."
            )
            text = extract_clean_text_with_anthropic(image, page_number, model)
            extracted_text.append({
                "page": page_number,
                "text": text or "",
            })

        full_text = "\\n\\n".join(page["text"] for page in extracted_text)

        if output_format == "txt":
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], f'anthropic_{job_id or uuid.uuid4().hex}.txt')
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(full_text)
            preview_text = full_text
        elif output_format == "docx":
            from docx import Document
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], f'anthropic_{job_id or uuid.uuid4().hex}.docx')
            doc = Document()
            doc.add_heading('Anthropic Extracted Text', 0)
            for page_data in extracted_text:
                doc.add_heading(f'Page {page_data["page"]}', level=1)
                for line_text in page_data["text"].splitlines():
                    doc.add_paragraph(line_text)
            doc.save(output_path)
            preview_text = full_text
        else:
            raise ValueError("Anthropic engine supports TXT, DOCX, and JSON output formats.")

    update_progress(100, "Done!")
    return {
        'text': preview_text,
        'pages': total_pages,
        'output_path': output_path
    }

"""

if "def process_with_anthropic" not in content:
    content = content.replace("def process_with_gemini(", anthropic_code + "\ndef process_with_gemini(")
    print("Added anthropic functions")

if "elif engine == 'anthropic':" not in content:
    content = content.replace(
        "elif engine == 'gemini':\n                result = process_with_gemini(file_path, output_format, job_id, page_range)",
        "elif engine == 'gemini':\n                result = process_with_gemini(file_path, output_format, job_id, page_range)\n            elif engine == 'anthropic':\n                result = process_with_anthropic(file_path, output_format, job_id, page_range)"
    )
    print("Patched process() router")

with open('app.py', 'w', encoding='utf-8') as f:
    f.write(content)

