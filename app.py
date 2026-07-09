import os
import uuid
import subprocess
import threading
from flask import Flask, render_template, request, jsonify, send_file, session, redirect, url_for
from werkzeug.utils import secure_filename
import fitz
from PIL import Image
import io
import easyocr
import numpy as np
import base64
import requests
import shutil
import re
import json
import mysql.connector
from mysql.connector import Error
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from flask_bcrypt import Bcrypt
from html import unescape
from jinja2 import utils

app = Flask(__name__)
app.secret_key = os.urandom(24)

bcrypt = Bcrypt(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'

class User(UserMixin):
    def __init__(self, id, username, name=None, email=None):
        self.id = id
        self.username = username
        self.name = name
        self.email = email

@login_manager.user_loader
def load_user(user_id):
    try:
        conn = mysql.connector.connect(**app.config['DB_CONFIG'])
        cursor = conn.cursor(dictionary=True)
        cursor.execute('SELECT * FROM admins WHERE id = %s', (user_id,))
        user_data = cursor.fetchone()
        cursor.close()
        conn.close()
        if user_data:
            return User(id=user_data['id'], username=user_data['username'],
                        name=user_data['name'], email=user_data['email'])
    except:
        pass
    return None

# Custom filter for JavaScript escaping
@app.template_filter('escapejs')
def escapejs_filter(s):
    if s is None:
        return ""
    # Standard JS escaping
    return str(s).replace('\\', '\\\\').replace("'", "\\'").replace('"', '\\"').replace('\n', '\\n').replace('\r', '\\r')

app.config['MAX_CONTENT_LENGTH'] = 1024 * 1024 * 1024
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'outputs'
app.config['DB_CONFIG'] = {
    'host': 'localhost',
    'user': 'job',
    'password': 'Xdman123456@',
    'database': 'job'
}

def get_db_setting(key, default=None):
    try:
        conn = mysql.connector.connect(**app.config['DB_CONFIG'])
        cursor = conn.cursor(dictionary=True)
        cursor.execute('SELECT `value` FROM settings WHERE `key` = %s', (key,))
        row = cursor.fetchone()
        cursor.close()
        conn.close()
        if row:
            return row['value']
    except:
        pass
    return default

# Global error handler to ensure JSON responses for AJAX calls
@app.errorhandler(Exception)
def handle_exception(e):
    # Pass through HTTP errors
    if hasattr(e, 'code') and isinstance(e.code, int):
        return jsonify(error=str(e), success=False), e.code
    # Handle non-HTTP exceptions only for AJAX requests
    if request.path.startswith('/admin') or request.path in ['/upload', '/process', '/engines', '/pipelines', '/progress']:
        return jsonify(error=str(e), success=False), 500
    # Otherwise return the default (which might be an HTML error page)
    return str(e), 500

# Global dictionary to track job progress
job_progress = {}
# Global dictionary to track job controls (threading.Event for pause/resume and cancel flag)
job_controls = {}

def parse_page_range(range_str, total_pages):
    if not range_str or range_str.lower() == 'all':
        return list(range(1, total_pages + 1))

    pages = set()
    parts = range_str.split(',')
    for part in parts:
        part = part.strip()
        if '-' in part:
            try:
                start, end = map(int, part.split('-'))
                for p in range(max(1, start), min(total_pages, end) + 1):
                    pages.add(p)
            except ValueError:
                continue
        else:
            try:
                p = int(part)
                if 1 <= p <= total_pages:
                    pages.add(p)
            except ValueError:
                continue
    return sorted(list(pages))

@app.route('/progress/<job_id>')
def get_progress(job_id):
    progress = job_progress.get(job_id, {"status": "unknown", "percentage": 0})
    return jsonify(progress)

@app.route('/stop/<job_id>', methods=['POST'])
def stop_job(job_id):
    if job_id in job_controls:
        job_controls[job_id]['pause_event'].clear()
        if job_id in job_progress:
            job_progress[job_id]['status'] = "Paused"
        return jsonify({'success': True})
    return jsonify({'error': 'Job not found'}), 404

@app.route('/resume/<job_id>', methods=['POST'])
def resume_job(job_id):
    if job_id in job_controls:
        job_controls[job_id]['pause_event'].set()
        if job_id in job_progress:
            job_progress[job_id]['status'] = "Resuming..."
        return jsonify({'success': True})
    return jsonify({'error': 'Job not found'}), 404

@app.route('/cancel/<job_id>', methods=['POST'])
def cancel_job(job_id):
    if job_id in job_controls:
        job_controls[job_id]['cancel_flag'] = True
        job_controls[job_id]['pause_event'].set() # Ensure it's not stuck in pause
        if job_id in job_progress:
            job_progress[job_id]['status'] = "Cancelled"
        return jsonify({'success': True})
    return jsonify({'error': 'Job not found'}), 404

def get_allowed_extensions():
    ext_str = get_db_setting('ALLOWED_EXTENSIONS', 'pdf,png,jpg,jpeg,bmp,tif,tiff,webp')
    return {ext.strip().lower() for ext in ext_str.split(',')}

def allowed_file(filename):
    allowed = get_allowed_extensions()
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

def init_db():
    try:
        conn = mysql.connector.connect(**app.config['DB_CONFIG'])
        cursor = conn.cursor()

        # Create tables based on normalized schema (MySQL syntax)
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS admins (
            id INT AUTO_INCREMENT PRIMARY KEY,
            username VARCHAR(100) UNIQUE NOT NULL,
            password VARCHAR(255) NOT NULL,
            name VARCHAR(255),
            email VARCHAR(255) UNIQUE
        )''')

        cursor.execute('''
        CREATE TABLE IF NOT EXISTS settings (
            id INT AUTO_INCREMENT PRIMARY KEY,
            `key` VARCHAR(100) UNIQUE NOT NULL,
            `value` TEXT,
            `description` TEXT,
            `category` VARCHAR(50) DEFAULT 'general'
        )''')

        # Insert default settings if they don't exist
        default_settings = [
            ('OPENAI_COMPAT_BASE_URL', 'http://localhost:8080/v1', 'Base URL for OpenAI-compatible API', 'openai'),
            ('OPENAI_COMPAT_API_KEY', 'pwd', 'API Key for OpenAI-compatible API', 'openai'),
            ('OPENAI_COMPAT_MODEL', 'gpt-5.5', 'Primary model name for OpenAI-compatible API', 'openai'),
            ('OPENAI_COMPAT_TIMEOUT_SECONDS', '240', 'Timeout for OpenAI-compatible API requests (seconds)', 'openai'),
            ('OPENAI_COMPAT_ENABLE_POSTPROCESS', 'false', 'Enable AI post-processing for local OCR', 'openai'),
            ('AISTUDIO_TOKEN', 'eabce52d47f2eacb24c9335a1a0e6b195e335efe', 'Token for AI Studio Layout API', 'aistudio'),
            ('AISTUDIO_API_URL', 'https://c7h8c1o6l62ej1ze.aistudio-app.com/layout-parsing', 'API URL for AI Studio Layout Parsing', 'aistudio'),
            ('AISTUDIO_TIMEOUT_SECONDS', '600', 'Timeout for AI Studio requests (seconds)', 'aistudio'),
            ('AISTUDIO_MAX_GEMINI_IMAGES', '80', 'Maximum images to send to Gemini per page', 'aistudio'),
            ('ANTHROPIC_API_KEY', '', 'API Key for Anthropic Claude models', 'anthropic'),
            ('ANTHROPIC_BASE_URL', '', 'Base URL for Anthropic-compatible API (leave empty for default)', 'anthropic'),
            ('ANTHROPIC_MODEL', 'claude-3-5-sonnet-latest', 'Primary model for Anthropic OCR', 'anthropic'),
            ('ANTHROPIC_TIMEOUT_SECONDS', '240', 'Timeout for Anthropic requests (seconds)', 'anthropic'),
            ('TESSERACT_PATH', '', 'Custom path to Tesseract executable (leave empty for auto-detect)', 'local_ocr'),
            ('EASYOCR_GPU', 'true', 'Use GPU for EasyOCR if available (true/false)', 'local_ocr'),
            ('OCR_DPI', '300', 'DPI for converting PDF pages to images (higher is more accurate but slower)', 'general'),
            ('OCR_MAX_RETRIES', '3', 'Maximum number of retries for API requests', 'general'),
            ('MAX_UPLOAD_SIZE_MB', '1024', 'Maximum allowed upload size in Megabytes', 'general'),
            ('ALLOWED_EXTENSIONS', 'pdf,png,jpg,jpeg,bmp,tif,tiff,webp', 'Comma-separated list of allowed file extensions', 'general'),
            ('FLASK_SECRET_KEY', 'super-secret-default-key-change-me', 'Secret key for Flask sessions (stay logged in across restarts)', 'general')
        ]

        for key, val, desc, cat in default_settings:
            cursor.execute('INSERT IGNORE INTO settings (`key`, `value`, `description`, `category`) VALUES (%s, %s, %s, %s)',
                         (key, val, desc, cat))

        cursor.execute('''
        CREATE TABLE IF NOT EXISTS ai_providers (
            id INT AUTO_INCREMENT PRIMARY KEY,
            name VARCHAR(100) NOT NULL,
            type ENUM('openai', 'gemini', 'ollama', 'aistudio', 'local', 'anthropic') NOT NULL,
            base_url VARCHAR(255),
            api_key VARCHAR(255),
            is_active BOOLEAN DEFAULT TRUE
        )''')

        cursor.execute('''
        CREATE TABLE IF NOT EXISTS custom_pipelines (
            id INT AUTO_INCREMENT PRIMARY KEY,
            name VARCHAR(100) NOT NULL,
            ocr_provider_id INT,
            ocr_model VARCHAR(100),
            structure_provider_id INT,
            structure_model VARCHAR(100),
            is_active BOOLEAN DEFAULT TRUE,
            FOREIGN KEY (ocr_provider_id) REFERENCES ai_providers(id) ON DELETE SET NULL,
            FOREIGN KEY (structure_provider_id) REFERENCES ai_providers(id) ON DELETE SET NULL
        )''')

        # Insert default providers if they don't exist
        cursor.execute("SELECT COUNT(*) FROM ai_providers")
        if cursor.fetchone()[0] == 0:
            default_providers = [
                ('Local OCR (Tesseract/EasyOCR)', 'local', '', ''),
                ('AI Studio (Paddle)', 'aistudio', 'https://c7h8c1o6l62ej1ze.aistudio-app.com/layout-parsing', 'eabce52d47f2eacb24c9335a1a0e6b195e335efe'),
                ('OpenAI API', 'openai', 'https://api.openai.com/v1', ''),
                ('Gemini API', 'gemini', '', ''),
                ('Anthropic API', 'anthropic', 'https://api.anthropic.com', ''),
                ('Ollama (Local)', 'ollama', 'http://localhost:11434/v1', '')
            ]
            for name, p_type, url, key in default_providers:
                cursor.execute('INSERT INTO ai_providers (name, type, base_url, api_key) VALUES (%s, %s, %s, %s)',
                             (name, p_type, url, key))

        # Insert default engines if they don't exist
        default_engines = [
            ('easyocr', 'EasyOCR (Local - Basic OCR)', 'Basic local OCR for English and Bengali.', True, 1),
            ('tesseract', 'Tesseract (Local - Fast OCR)', 'Fast local OCR using Tesseract.', True, 2),
            ('openai_compatible', 'OpenAI Compatible (Local API - Vision OCR)', 'Vision-based OCR using an OpenAI-compatible API.', True, 3),
            ('gemini', 'Gemini (Local Proxy - OCR + Cleanup + Structure)', 'Direct page-to-JSON extraction using Gemini.', True, 4),
            ('aistudio', 'AI Studio Layout API (Cloud - Layout & Images)', 'Advanced cloud layout parsing and image extraction.', True, 5),
            ('anthropic', 'Anthropic Claude (Cloud - Vision OCR & Structure)', 'Vision-based OCR using Anthropic Claude models.', True, 6),
        ]

        for name, display, desc, active, order in default_engines:
            cursor.execute('INSERT IGNORE INTO processing_engines (name, display_name, description, is_active, sort_order) VALUES (%s, %s, %s, %s, %s)',
                         (name, display, desc, active, order))

        cursor.execute('''
        CREATE TABLE IF NOT EXISTS exams (
            id INT AUTO_INCREMENT PRIMARY KEY,
            name VARCHAR(255) UNIQUE NOT NULL
        )''')

        cursor.execute('''
        CREATE TABLE IF NOT EXISTS years (
            id INT AUTO_INCREMENT PRIMARY KEY,
            year VARCHAR(100) UNIQUE NOT NULL
        )''')

        cursor.execute('''
        CREATE TABLE IF NOT EXISTS questions (
            id INT AUTO_INCREMENT PRIMARY KEY,
            text TEXT NOT NULL
        )''')

        cursor.execute('''
        CREATE TABLE IF NOT EXISTS options (
            id INT AUTO_INCREMENT PRIMARY KEY,
            text TEXT,
            type VARCHAR(50) DEFAULT 'text',
            image_base64 LONGTEXT
        )''')

        cursor.execute('''
        CREATE TABLE IF NOT EXISTS mcqs (
            id INT AUTO_INCREMENT PRIMARY KEY,
            question_id INT,
            option_ids TEXT, -- JSON string of option IDs
            answer_index INT, -- 1-4
            answer_id INT, -- FK to options.id
            explanation TEXT,
            language VARCHAR(50),
            subject VARCHAR(255),
            FOREIGN KEY (question_id) REFERENCES questions(id),
            FOREIGN KEY (answer_id) REFERENCES options(id)
        )''')

        cursor.execute('''
        CREATE TABLE IF NOT EXISTS exam_questions (
            id INT AUTO_INCREMENT PRIMARY KEY,
            exam_id INT,
            year_id INT,
            mcq_id INT,
            FOREIGN KEY (exam_id) REFERENCES exams(id),
            FOREIGN KEY (year_id) REFERENCES years(id),
            FOREIGN KEY (mcq_id) REFERENCES mcqs(id)
        )''')

        conn.commit()
        cursor.close()
        conn.close()
        print("MySQL Database initialized successfully.")
    except Error as e:
        print(f"Error while connecting to MySQL: {e}")

def cleanup_temp_files():
    """Remove legacy temp files and directories on startup."""
    for folder in [app.config['UPLOAD_FOLDER'], app.config['OUTPUT_FOLDER']]:
        if not os.path.exists(folder):
            continue
        for item in os.listdir(folder):
            if item.startswith('temp_'):
                item_path = os.path.join(folder, item)
                try:
                    if os.path.isdir(item_path):
                        shutil.rmtree(item_path)
                    else:
                        os.remove(item_path)
                except Exception as e:
                    print(f"Failed to cleanup {item_path}: {e}")

def get_provider_details(provider_id):
    if not provider_id:
        return None
    try:
        conn = mysql.connector.connect(**app.config['DB_CONFIG'])
        cursor = conn.cursor(dictionary=True)
        cursor.execute('SELECT * FROM ai_providers WHERE id = %s', (provider_id,))
        provider = cursor.fetchone()
        cursor.close()
        conn.close()
        return provider
    except:
        return None

def run_custom_pipeline(file_path, pipeline_id, output_format, job_id=None, page_range=None, language='english'):
    config = get_config()
    def update_progress(percentage, status):
        if job_id:
            job_progress[job_id] = {"status": status, "percentage": percentage}

    def check_controls():
        if job_id and job_id in job_controls:
            if job_controls[job_id].get('cancel_flag'):
                raise InterruptedError("Job cancelled by user")
            job_controls[job_id]['pause_event'].wait()

    # 1. Fetch Pipeline Details
    try:
        conn = mysql.connector.connect(**app.config['DB_CONFIG'])
        cursor = conn.cursor(dictionary=True)
        cursor.execute('SELECT * FROM custom_pipelines WHERE id = %s', (pipeline_id,))
        pipeline = cursor.fetchone()
        cursor.close()
        conn.close()
    except Error as e:
        raise RuntimeError(f"Database error: {e}")

    if not pipeline:
        raise ValueError("Pipeline not found")

    ocr_provider = get_provider_details(pipeline['ocr_provider_id'])
    struct_provider = get_provider_details(pipeline['structure_provider_id'])

    # 2. Convert PDF to Images
    total_pdf_pages = 0
    if is_pdf_file(file_path):
        doc = fitz.open(file_path)
        total_pdf_pages = len(doc)
        doc.close()

    page_indices = parse_page_range(page_range, total_pdf_pages) if is_pdf_file(file_path) else None
    page_images = get_document_page_images(file_path, dpi=config['OCR_DPI'], page_indices=page_indices)
    total_pages = len(page_images)

    extracted_pages = []

    # 3. Process each page
    for idx, (page_num, image) in enumerate(page_images):
        check_controls()
        update_progress(5 + int((idx / total_pages) * 45), f"OCR Processing page {page_num} of {total_pages}...")

        # Step 1: OCR
        page_text = perform_ocr_with_provider(ocr_provider, pipeline['ocr_model'], image, page_num, language)

        extracted_pages.append({
            "page": page_num,
            "text": page_text
        })

    full_text = "\n\n".join(p['text'] for p in extracted_pages)

    # 4. Step 2: Structuring (If output format is JSON)
    if output_format == 'json':
        update_progress(60, "Structuring text into JSON dataset...")
        dataset = structure_text_with_provider(struct_provider, pipeline['structure_model'], full_text)

        output_path = os.path.join(app.config['OUTPUT_FOLDER'], f'pipeline_{job_id or uuid.uuid4().hex}.json')
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(dataset, f, ensure_ascii=False, indent=2)
        preview_text = json.dumps(dataset, ensure_ascii=False, indent=2)
    else:
        # Standard text/docx output
        if output_format == 'txt':
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], f'pipeline_{job_id or uuid.uuid4().hex}.txt')
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(full_text)
            preview_text = full_text
        elif output_format == 'docx':
            from docx import Document
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], f'pipeline_{job_id or uuid.uuid4().hex}.docx')
            doc = Document()
            doc.add_heading('Pipeline Extracted Text', 0)
            for page_data in extracted_pages:
                doc.add_heading(f'Page {page_data["page"]}', level=1)
                for line in page_data["text"].splitlines():
                    doc.add_paragraph(line)
            doc.save(output_path)
            preview_text = full_text
        else:
            raise ValueError("Unsupported format")

    update_progress(100, "Pipeline Completed!")
    return {
        'text': preview_text,
        'pages': total_pages,
        'output_path': output_path
    }

def perform_ocr_with_provider(provider, model, image, page_num, language='english'):
    if not provider:
        return "[Error: No OCR Provider selected]"

    p_type = provider['type']

    if p_type == 'local':
        # Default to EasyOCR if model is not specified or 'easyocr'
        if not model or 'easyocr' in model.lower():
            import numpy as np
            img_np = np.array(image)
            langs = get_easyocr_langs(language)
            reader = get_reader(langs)
            result = reader.readtext(img_np, detail=1)
            return reconstruct_layout(result)
        else:
            # Fallback to Tesseract
            import pytesseract
            configure_tesseract_command(pytesseract)
            lang_map = {'english': 'eng', 'bengali': 'ben', 'both': 'eng+ben'}
            tess_lang = lang_map.get(language, 'eng')
            return pytesseract.image_to_string(image, lang=tess_lang)

    elif p_type == 'aistudio':
        # AI Studio uses its own specific logic
        # For simplicity in custom pipeline, we'll use a simplified version of the existing logic
        # but this would usually be its own flow.
        return extract_clean_text_with_gemini(image, page_num, model or 'gemini-3-flash')

    elif p_type in ['openai', 'ollama', 'gemini', 'anthropic']:
        # Generic Vision API call
        data_url = prepare_openai_compatible_image_data_url(image)
        messages = [{
            "role": "user",
            "content": [
                {"type": "text", "text": "Extract all text from this image. Return only the extracted text."},
                {"type": "image_url", "image_url": {"url": data_url}}
            ]
        }]

        if p_type == 'openai':
            response = call_custom_openai_api(provider['base_url'], provider['api_key'], model or 'gpt-4o', messages)
            return response
        elif p_type == 'anthropic':
            return call_anthropic_api(model or 'claude-3-5-sonnet-latest', messages)
        elif p_type == 'gemini':
            # Use the existing gemini logic but with the provider key if available
            return extract_clean_text_with_gemini(image, page_num, model or 'gemini-3-flash')
        elif p_type == 'ollama':
            return call_ollama_vision(provider['base_url'], model or 'llava', data_url)

    return f"[Error: Provider type {p_type} not implemented for OCR]"

def structure_text_with_provider(provider, model, text):
    if not provider:
        return {"mcqs": [], "error": "No structuring provider"}

    messages = [
        {"role": "system", "content": MCQ_DATASET_SYSTEM_PROMPT},
        {"role": "user", "content": f"Convert the following OCR text into JSON:\n\n{text}"}
    ]

    p_type = provider['type']

    try:
        if p_type == 'openai':
            content = call_custom_openai_api(provider['base_url'], provider['api_key'], model or 'gpt-4o', messages)
        elif p_type == 'anthropic':
            content = call_anthropic_api(model or 'claude-3-5-sonnet-latest', messages)
        elif p_type == 'gemini':
            # Simplified: using existing logic
            caps = get_openai_compat_capabilities()
            content = get_openai_compatible_message_text(call_openai_compatible_chat(model or caps['text_model'], messages))
        elif p_type == 'ollama':
            content = call_ollama_chat(provider['base_url'], model or 'llama3', messages)
        else:
            # Fallback to default Gemini structuring
            return transform_ocr_text_to_mcq_dataset(text, get_config()['OPENAI_COMPAT_MODEL'])

        return normalize_mcq_dataset(parse_json_llm_response(content))
    except Exception as e:
        return {"mcqs": [], "error": str(e)}

def call_custom_openai_api(base_url, api_key, model, messages):
    from openai import OpenAI
    client = OpenAI(base_url=base_url or "https://api.openai.com/v1", api_key=api_key or "sk-...")
    response = client.chat.completions.create(
        model=model,
        messages=messages,
        temperature=0
    )
    return response.choices[0].message.content

def call_ollama_chat(base_url, model, messages):
    # Standardize Ollama URL to point to /api if not specified
    if not base_url:
        base_url = "http://localhost:11434"

    url = f"{base_url.rstrip('/')}/api/chat"
    if '/v1' in base_url:
        # If user provided a v1 URL, they probably want OpenAI compatibility
        # We can either handle it here or tell them to use the 'openai' provider type
        url = f"{base_url.rstrip('/')}/chat/completions"
        # ... actually, better to use native Ollama API for 'ollama' type
        url = base_url.replace('/v1', '/api/chat')

    payload = {
        "model": model or "llama3",
        "messages": messages,
        "stream": False,
        "options": {"temperature": 0}
    }
    response = requests.post(url, json=payload, timeout=120)
    response.raise_for_status()
    return response.json()['message']['content']

def call_ollama_vision(base_url, model, data_url):
    if not base_url:
        base_url = "http://localhost:11434"

    url = f"{base_url.rstrip('/')}/api/generate"
    if '/v1' in base_url:
        url = base_url.replace('/v1', '/api/generate')

    base64_data = data_url.split(',')[1]
    payload = {
        "model": model or "llava",
        "prompt": "Extract all text from this image. Return only the text.",
        "images": [base64_data],
        "stream": False
    }
    response = requests.post(url, json=payload, timeout=120)
    response.raise_for_status()
    return response.json()['response']

def get_config():
    return {
        'OPENAI_COMPAT_BASE_URL': get_db_setting('OPENAI_COMPAT_BASE_URL', os.environ.get("OPENAI_COMPAT_BASE_URL", "http://localhost:8045/v1")),
        'OPENAI_COMPAT_API_KEY': get_db_setting('OPENAI_COMPAT_API_KEY', os.environ.get("OPENAI_COMPAT_API_KEY", "sk-28d07728e1aa4ac5adb0d1fc09b7d743")),
        'OPENAI_COMPAT_MODEL': get_db_setting('OPENAI_COMPAT_MODEL', os.environ.get("OPENAI_COMPAT_MODEL", "gemini-3-flash")),
        'OPENAI_COMPAT_TIMEOUT_SECONDS': float(get_db_setting('OPENAI_COMPAT_TIMEOUT_SECONDS', "240")),
        'AISTUDIO_TOKEN': get_db_setting('AISTUDIO_TOKEN', "eabce52d47f2eacb24c9335a1a0e6b195e335efe"),
        'AISTUDIO_API_URL': get_db_setting('AISTUDIO_API_URL', "https://c7h8c1o6l62ej1ze.aistudio-app.com/layout-parsing"),
        'AISTUDIO_TIMEOUT_SECONDS': float(get_db_setting('AISTUDIO_TIMEOUT_SECONDS', "600")),
        'AISTUDIO_MAX_GEMINI_IMAGES': int(get_db_setting('AISTUDIO_MAX_GEMINI_IMAGES', "80")),
        'OPENAI_COMPAT_ENABLE_POSTPROCESS': get_db_setting('OPENAI_COMPAT_ENABLE_POSTPROCESS', "false").lower() == "true",
        'TESSERACT_PATH': get_db_setting('TESSERACT_PATH', ""),
        'EASYOCR_GPU': get_db_setting('EASYOCR_GPU', "true").lower() == "true",
        'OCR_DPI': int(get_db_setting('OCR_DPI', "300")),
        'OCR_MAX_RETRIES': int(get_db_setting('OCR_MAX_RETRIES', "3")),
        'ANTHROPIC_API_KEY': get_db_setting('ANTHROPIC_API_KEY', ""),
        'ANTHROPIC_BASE_URL': get_db_setting('ANTHROPIC_BASE_URL', ""),
        'ANTHROPIC_MODEL': get_db_setting('ANTHROPIC_MODEL', "claude-3-5-sonnet-latest"),
        'ANTHROPIC_TIMEOUT_SECONDS': float(get_db_setting('ANTHROPIC_TIMEOUT_SECONDS', "240"))
    }

# Original global variables (will now be accessed via get_config() where needed or updated periodically)
# For better performance, we can wrap these in a function or a class.

OPENAI_COMPAT_MODEL_FALLBACKS = os.environ.get("OPENAI_COMPAT_MODEL_FALLBACKS", "")
OPENAI_COMPAT_TIMEOUT_SECONDS = float(os.environ.get("OPENAI_COMPAT_TIMEOUT_SECONDS", "240"))
OPENAI_COMPAT_PROBE_TIMEOUT_SECONDS = float(os.environ.get("OPENAI_COMPAT_PROBE_TIMEOUT_SECONDS", "12"))

# Initialize DB and cleanup on startup
init_db()
cleanup_temp_files()

# Set persistent secret key from DB
persistent_key = get_db_setting('FLASK_SECRET_KEY')
if persistent_key:
    app.secret_key = persistent_key.encode('utf-8')

MCQ_DATASET_SYSTEM_PROMPT = """
You are an AI data processing agent designed to convert OCR-extracted text from Bangla, English, or mixed-language government job preparation books into structured, high-quality datasets.

Objective:
Transform noisy OCR text into:
1. Structured MCQ data for database storage
2. Clean contextual text for AI retrieval / RAG systems

Text cleaning rules:
- Normalize whitespace.
- Fix broken words only when the correction is clear from context.
- Remove irrelevant symbols, headers, page numbers, and noise.
- Preserve Bangla characters correctly (Unicode range U+0980-U+09FF).
- Keep semantic meaning intact.

MCQ extraction rules:
- Extract all complete MCQs from the cleaned text.
- Each MCQ must contain question, exactly 4 options, answer, explanation, language, subject, exam, and year.
- Options may be marked A/B/C/D or ক/খ/গ/ঘ.
- Text-only options may be returned as strings.
- When an option is represented by an image, return an object with type "image", text null, and imageBase64 set to the provided data URL.
- When an option has both visible text and an associated image, return an object with type "image", text set to the visible label/text, and imageBase64 set to the provided data URL.
- answer must be the correct option number as an integer: 1, 2, 3, or 4.
- If an answer/explanation is explicitly present but does not match any of the 4 options, use 0.
- If no answer evidence is explicitly detectable, use null.
- If the explanation is not available, use null.
- If options are unclear, skip that MCQ.
- For AI Studio page-image inputs, do not omit visible numbered questions. Use the page image to recover answers/options that are missing from markdown text.
- Do not hallucinate answers, explanations, subjects, exams, years, or missing options.
- Prefer high accuracy over completeness.
- Do not merge multiple questions incorrectly.
- language must be exactly "bn", "en", or "mixed".

Context generation rules:
- Generate one clean declarative knowledge paragraph for each MCQ when enough facts are known.
- Convert MCQs into factual statements.
- Keep both Bangla and English when present.
- If the correct answer is null and a factual context cannot be determined safely, omit that context.

Return ONLY valid JSON. Do not include markdown, comments, or explanation.
The JSON object must have this exact top-level structure:
{
  "mcqs": [
    {
      "question": "",
      "options": [
        {"type": "text", "text": "", "imageBase64": null},
        {"type": "text", "text": "", "imageBase64": null},
        {"type": "text", "text": "", "imageBase64": null},
        {"type": "image", "text": null, "imageBase64": "data:image/jpeg;base64,..."}
      ],
      "answer": 1,
      "explanation": null,
      "language": "bn",
      "subject": null,
      "exam": null,
      "year": null
    }
  ],
  "contexts": []
}
""".strip()


def allowed_file(filename):
    allowed = get_allowed_extensions()
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed

IMAGE_EXTENSIONS = {'png', 'jpg', 'jpeg', 'bmp', 'tif', 'tiff', 'webp'}

def get_file_extension(path):
    return os.path.splitext(path)[1].lower().lstrip('.')


def is_pdf_file(path):
    return get_file_extension(path) == 'pdf'


def is_image_file(path):
    return get_file_extension(path) in IMAGE_EXTENSIONS


def get_easyocr_langs(lang):
    mapping = {
        'english': ['en'],
        'bengali': ['bn'],
        'both': ['en', 'bn']
    }
    return mapping.get(lang, ['en'])


# Cache for EasyOCR readers to avoid re-initializing
OcrReaders = {}
OpenAICompatClient = None
OpenAICompatCapabilities = None
OpenAICompatProbeErrors = []
OpenAICompatModelListFailed = False

def get_reader(langs):
    lang_key = tuple(sorted(langs))
    config = get_config()
    if lang_key not in OcrReaders:
        # Initialize reader (this downloads models if first time)
        OcrReaders[lang_key] = easyocr.Reader(list(langs), gpu=config['EASYOCR_GPU'])
    return OcrReaders[lang_key]


def get_openai_compat_client():
    global OpenAICompatClient
    config = get_config()

    if OpenAICompatClient is None:
        try:
            from openai import OpenAI
        except ImportError as exc:
            raise RuntimeError(
                "The OpenAI Python package is required for the OpenAI-compatible engine. "
                "Install it with: pip install openai"
            ) from exc

        OpenAICompatClient = OpenAI(
            base_url=config['OPENAI_COMPAT_BASE_URL'],
            api_key=config['OPENAI_COMPAT_API_KEY'],
            timeout=config['OPENAI_COMPAT_TIMEOUT_SECONDS'],
            max_retries=0,
        )

    return OpenAICompatClient


def call_openai_compatible_chat(model, messages, temperature=0, timeout=None):
    client = get_openai_compat_client()
    config = get_config()
    return client.chat.completions.create(
        model=model,
        messages=messages,
        temperature=temperature,
        timeout=timeout or config['OPENAI_COMPAT_TIMEOUT_SECONDS'],
    )


def list_openai_compatible_models():
    global OpenAICompatModelListFailed
    config = get_config()

    client = get_openai_compat_client()

    try:
        response = client.models.list()
        OpenAICompatModelListFailed = False
        return [model.id for model in response.data if getattr(model, "id", None)]
    except Exception as exc:
        OpenAICompatModelListFailed = True
        OpenAICompatProbeErrors.append(f"Could not list models from {config['OPENAI_COMPAT_BASE_URL']}: {exc}")
        return []


def get_openai_compatible_model_candidates():
    discovered = list_openai_compatible_models()
    preferred = []
    config = get_config()

    if config['OPENAI_COMPAT_MODEL']:
        preferred.append(config['OPENAI_COMPAT_MODEL'])

    preferred.extend(
        model.strip()
        for model in OPENAI_COMPAT_MODEL_FALLBACKS.split(",")
        if model.strip()
    )

    if discovered or OPENAI_COMPAT_MODEL_FALLBACKS:
        preferred.extend(
            [
                "gemini-3-flash",
                "gemini-3-pro",
                "gemini-3-pro-high",
                "gemini-3-pro-low",
                "gemini-3-pro-preview",
                "gemini-2.5-flash",
                "gemini-2.5-pro",
                "gemini-2.0-flash",
                "gpt-4o-mini",
                "gpt-4o",
            ]
        )

    candidates = []
    seen = set()
    for model in preferred + discovered:
        if not model or model in seen:
            continue
        seen.add(model)
        candidates.append(model)

    return candidates


def get_openai_compatible_message_text(response):
    message = response.choices[0].message.content
    if isinstance(message, str):
        return message.strip()

    if isinstance(message, list):
        parts = []
        for item in message:
            if isinstance(item, dict):
                if item.get("type") == "text":
                    parts.append(item.get("text", ""))
                elif "text" in item:
                    parts.append(item.get("text", ""))
        return "\n".join(part.strip() for part in parts if part and part.strip())

    return ""


def supports_openai_text_model(model):
    try:
        response = call_openai_compatible_chat(
            model,
            [{"role": "user", "content": "Reply with exactly OK"}],
            temperature=0,
            timeout=OPENAI_COMPAT_PROBE_TIMEOUT_SECONDS,
        )
        content = get_openai_compatible_message_text(response)
        if content:
            return True
        OpenAICompatProbeErrors.append(f"Model {model} returned an empty text response.")
        return False
    except Exception as exc:
        OpenAICompatProbeErrors.append(f"Model {model} text probe failed: {exc}")
        return False


def supports_openai_vision_model(model):
    tiny_image = "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+aS1cAAAAASUVORK5CYII="

    try:
        response = call_openai_compatible_chat(
            model,
            [
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "What is in this image? Reply with one short sentence."},
                        {
                            "type": "image_url",
                            "image_url": {"url": f"data:image/png;base64,{tiny_image}"},
                        },
                    ],
                }
            ],
            temperature=0,
            timeout=OPENAI_COMPAT_PROBE_TIMEOUT_SECONDS,
        )
        content = get_openai_compatible_message_text(response)
        return bool(content)
    except Exception:
        return False


def get_openai_compat_capabilities():
    global OpenAICompatCapabilities
    config = get_config()

    if OpenAICompatCapabilities is not None:
        return OpenAICompatCapabilities

    OpenAICompatProbeErrors.clear()
    candidates = get_openai_compatible_model_candidates()
    text_model = None

    for model in candidates:
        if supports_openai_text_model(model):
            text_model = model
            break

    if not text_model:
        details = " ".join(OpenAICompatProbeErrors[-4:]).strip()
        configured = config['OPENAI_COMPAT_MODEL'] or "(not set)"
        suffix = f" Details: {details}" if details else ""
        raise RuntimeError(
            "No working text model was found on the OpenAI-compatible endpoint. "
            f"Base URL: {config['OPENAI_COMPAT_BASE_URL']}. Configured model: {configured}. "
            "Start the proxy service or set OPENAI_COMPAT_BASE_URL / OPENAI_COMPAT_MODEL "
            "to a model that supports chat completions."
            f"{suffix}"
        )

    vision_model = text_model if supports_openai_vision_model(text_model) else None

    OpenAICompatCapabilities = {
        "text_model": text_model,
        "vision_model": vision_model,
        "vision_supported": vision_model is not None,
    }
    return OpenAICompatCapabilities


def prepare_openai_compatible_image_data_url(image):
    # Keep request bodies small enough for local OpenAI-compatible servers.
    max_payload_bytes = 3_500_000
    attempts = [
        (1800, 85),
        (1400, 75),
        (1100, 65),
        (900, 55),
    ]

    source = image.convert("RGB")

    for max_edge, quality in attempts:
        candidate = source.copy()
        candidate.thumbnail((max_edge, max_edge), Image.Resampling.LANCZOS)

        buffer = io.BytesIO()
        candidate.save(buffer, format="JPEG", quality=quality, optimize=True)
        image_bytes = buffer.getvalue()

        if len(image_bytes) <= max_payload_bytes:
            encoded = base64.b64encode(image_bytes).decode("ascii")
            return f"data:image/jpeg;base64,{encoded}"

    # Fallback to the smallest attempt even if the server limit is unusually low.
    buffer = io.BytesIO()
    source.thumbnail((700, 700), Image.Resampling.LANCZOS)
    source.save(buffer, format="JPEG", quality=45, optimize=True)
    encoded = base64.b64encode(buffer.getvalue()).decode("ascii")
    return f"data:image/jpeg;base64,{encoded}"


def extract_text_with_openai_compatible(image, page_number, model):
    data_url = prepare_openai_compatible_image_data_url(image)

    response = call_openai_compatible_chat(
        model,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            "You are an OCR engine. Extract all visible text from this document page. "
                            "Preserve line breaks and reading order as well as possible. "
                            "Return only the extracted text."
                        ),
                    },
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": data_url,
                        },
                    },
                ],
            }
        ],
        temperature=0,
    )

    message = get_openai_compatible_message_text(response)
    if message:
        return message

    raise RuntimeError(f"No text returned by OpenAI-compatible engine for page {page_number}.")


def get_document_page_images(file_path, dpi=200, page_indices=None):
    page_images = []

    if is_pdf_file(file_path):
        doc = fitz.open(file_path)
        try:
            total_pdf_pages = len(doc)
            if page_indices is None:
                indices = range(total_pdf_pages)
            else:
                # Convert 1-based to 0-based
                indices = [i-1 for i in page_indices if 0 <= i-1 < total_pdf_pages]

            for page_num in indices:
                page = doc[page_num]
                pix = page.get_pixmap(dpi=dpi)
                img_bytes = pix.tobytes("png")
                img = Image.open(io.BytesIO(img_bytes))
                page_images.append((page_num + 1, img))
        finally:
            doc.close()
    elif is_image_file(file_path):
        img = Image.open(file_path)
        page_images.append((1, img))
    else:
        raise ValueError("Unsupported file type. Please upload a PDF or image file.")

    return page_images


def extract_clean_text_with_gemini(image, page_number, model):
    data_url = prepare_openai_compatible_image_data_url(image)

    response = call_openai_compatible_chat(
        model,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            "OCR this document page and clean the result. Preserve Bangla and English exactly. "
                            "Keep question numbers, options, formulas, answer markers such as 'উ. ক', and explanations. "
                            "Do not summarize. Return only cleaned page text."
                        ),
                    },
                    {"type": "image_url", "image_url": {"url": data_url}},
                ],
            }
        ],
        temperature=0,
    )

    message = get_openai_compatible_message_text(response)
    if message:
        return message

    raise RuntimeError(f"No text returned by Gemini engine for page {page_number}.")


def transform_page_image_to_mcq_dataset_with_gemini(image, page_number, model, page_text=None):
    data_url = prepare_openai_compatible_image_data_url(image)

    response = call_openai_compatible_chat(
        model,
        messages=[
            {"role": "system", "content": MCQ_DATASET_SYSTEM_PROMPT},
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            f"OCR, clean, and structure all complete MCQs visible on page {page_number}. "
                            "Use both the cleaned OCR text below and the page image directly for OCR, options, formulas, "
                            "answer markers, and explanations. "
                            "Do not omit visible numbered questions. Right-side printed markers like 'উ. ক', 'উ. খ', "
                            "'উ. গ', and 'উ. ঘ' map to answer 1, 2, 3, and 4. "
                            "If an explanation gives an answer that is not one of the four options, set answer to 0. "
                            "If no answer evidence is visible, set answer to null. "
                            "Include image/pattern/diagram questions too; if option images cannot be cropped, describe each option image briefly in text. "
                            "Return only strict JSON.\n\n"
                            "Cleaned OCR text from this same page:\n"
                            f"{page_text or ''}"
                        ),
                    },
                    {"type": "image_url", "image_url": {"url": data_url}},
                ],
            },
        ],
        temperature=0,
    )

    message = get_openai_compatible_message_text(response)
    try:
        parsed = parse_json_llm_response(message)
    except json.JSONDecodeError as exc:
        raise ValueError(f"Gemini structured extraction returned invalid JSON on page {page_number}.") from exc

    return normalize_mcq_dataset(parsed)


def normalize_digit_text(value):
    digit_map = str.maketrans("০১২৩৪৫৬৭৮৯", "0123456789")
    return str(value).translate(digit_map)


def get_longest_consecutive_question_count(text):
    numbers = []
    for match in re.finditer(r"(?<![\d০-৯])([০-৯0-9]{1,3})\s*[\.)।]", text or ""):
        try:
            number = int(normalize_digit_text(match.group(1)))
        except ValueError:
            continue
        if 1 <= number <= 250:
            numbers.append(number)

    if not numbers:
        return 0

    unique_numbers = sorted(set(numbers))
    best = 1
    current = 1
    for previous, number in zip(unique_numbers, unique_numbers[1:]):
        if number == previous + 1:
            current += 1
            best = max(best, current)
        else:
            current = 1

    return best


def get_question_number_from_text(text):
    match = re.match(r"^\s*([০-৯0-9]{1,3})\s*[\.)।]", text or "")
    if not match:
        return None
    try:
        return int(normalize_digit_text(match.group(1)))
    except ValueError:
        return None


def extract_missing_mcqs_with_gemini(image, page_number, page_text, existing_dataset, model):
    data_url = prepare_openai_compatible_image_data_url(image)
    existing_questions = [
        mcq.get("question") or ""
        for mcq in existing_dataset.get("mcqs", [])
    ]

    response = call_openai_compatible_chat(
        model,
        messages=[
            {"role": "system", "content": MCQ_DATASET_SYSTEM_PROMPT},
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            f"On page {page_number}, some visible numbered MCQs were missed. "
                            "Extract ONLY complete MCQs visible on the page image that are NOT already in the previous list. "
                            "Pay special attention to image, diagram, series, projectile, pattern, and figure-based questions. "
                            "Use the cleaned OCR text and the page image. If an option is a diagram and cannot be cropped, "
                            "describe it briefly as text. Return only strict JSON.\n\n"
                            "Previously extracted questions:\n"
                            f"{json.dumps(existing_questions, ensure_ascii=False)}\n\n"
                            "Cleaned OCR text:\n"
                            f"{page_text}"
                        ),
                    },
                    {"type": "image_url", "image_url": {"url": data_url}},
                ],
            },
        ],
        temperature=0,
    )

    message = get_openai_compatible_message_text(response)
    try:
        parsed = parse_json_llm_response(message)
    except json.JSONDecodeError:
        return {"mcqs": [], "contexts": []}

    return normalize_mcq_dataset(parsed)


def should_fallback_openai_vision(exc):
    message = str(exc).lower()
    markers = [
        "unable to process input image",
        "unknown error",
        "server_error",
        "error code: 500",
    ]
    return any(marker in message for marker in markers)


def extract_text_with_tesseract_image(image, language):
    import pytesseract
    from pytesseract import Output

    if not configure_tesseract_command(pytesseract):
        raise RuntimeError(
            "Tesseract executable not found. Install Tesseract OCR or add tesseract.exe to your PATH."
        )

    lang_map = {'english': 'eng', 'bengali': 'ben', 'both': 'eng+ben'}
    tess_lang = lang_map.get(language, 'eng')

    # pytesseract can take PIL images directly
    try:
        data = pytesseract.image_to_data(image, lang=tess_lang, output_type=Output.DICT)
    except Exception as e:
        raise RuntimeError(f"Tesseract processing failed: {e}")

    ocr_results = []
    n_boxes = len(data['text'])
    for i in range(n_boxes):
        if int(data['conf'][i]) > -1:
            val = data['text'][i].strip()
            if val:
                x, y, w, h = data['left'][i], data['top'][i], data['width'][i], data['height'][i]
                bounds = [[x, y], [x+w, y], [x+w, y+h], [x, y+h]]
                ocr_results.append((bounds, val, data['conf'][i]))

    return reconstruct_layout(ocr_results)


def extract_text_with_easyocr_image(image, language):
    langs = get_easyocr_langs(language)
    reader = get_reader(langs)

    # EasyOCR can take NumPy arrays or file paths. PIL to NumPy is efficient.
    import numpy as np
    try:
        img_np = np.array(image.convert('RGB'))
        result = reader.readtext(img_np, detail=1)
    except Exception as e:
        raise RuntimeError(f"EasyOCR processing failed: {e}")

    return reconstruct_layout(result)


def extract_text_with_local_ocr_fallback(image, language):
    try:
        return extract_text_with_tesseract_image(image, language)
    except Exception:
        return extract_text_with_easyocr_image(image, language)


def cleanup_ocr_text_with_openai_compatible(raw_text, model):
    if not raw_text or not raw_text.strip():
        return ""

    response = call_openai_compatible_chat(
        model,
        messages=[
            {
                "role": "user",
                "content": (
                    "You are a conservative OCR post-processor. "
                    "Do not paraphrase. Do not summarize. Do not add any words that are not present. "
                    "Keep line breaks and order. "
                    "If uncertain, keep the original token unchanged. "
                    "Return only cleaned OCR text.\n\n"
                    f"{raw_text}"
                ),
            }
        ],
        temperature=0,
    )

    message = get_openai_compatible_message_text(response)
    if message:
        return message

    return raw_text


def strip_json_code_fence(content):
    text = (content or "").strip()
    
    # Try to find a JSON block between triple backticks
    import re
    match = re.search(r'```(?:json)?\s*(.*?)\s*```', text, flags=re.IGNORECASE | re.DOTALL)
    if match:
        return match.group(1).strip()
        
    # If no backticks, try to find the first '{' and the last '}'
    start = text.find('{')
    end = text.rfind('}')
    if start != -1 and end != -1 and end > start:
        return text[start:end+1].strip()
        
    # If all else fails, return the original text
    return text



def parse_json_llm_response(message):
    import ast
    import re
    import json
    import logging
    text = strip_json_code_fence(message)
    try:
        return json.loads(text)
    except Exception as e:
        # Fallback 1: Fix trailing commas
        try:
            text_no_trailing = re.sub(r',\s*([\]}])', r'\1', text)
            return json.loads(text_no_trailing)
        except Exception:
            pass
            
        # Fallback 2: Try ast.literal_eval in case LLM output a Python dictionary with single quotes
        try:
            if text.startswith('{') or text.startswith('['):
                return ast.literal_eval(text)
        except Exception:
            pass
            
        # Fallback 3: Return empty dataset instead of crashing
        logging.warning(f"Failed to parse LLM JSON. Returning empty dataset. Raw text: {text[:200]}...")
        return {"mcqs": [], "contexts": []}


def extract_option_number_from_text(text, options):
    if not text:
        return None

    text_value = str(text)
    markers = [
        r"(?:correct\s+answer|answer|ans)\s*[:：\-]?\s*([^\n,;।]+)",
        r"(?:সঠিক\s+উত্তর|উত্তর)\s*[:：\-]?\s*([^\n,;।]+)",
    ]

    for pattern in markers:
        match = re.search(pattern, text_value, flags=re.IGNORECASE)
        if match:
            normalized = normalize_answer_to_option_number(match.group(1), options)
            if normalized is not None:
                return normalized

    return normalize_answer_to_option_number(text_value, options)


def contains_answer_evidence(text):
    if not text:
        return False

    text_value = str(text).strip()
    if not text_value:
        return False

    markers = [
        r"(?:correct\s+answer|answer|ans)\s*[:：\-]?\s*[^\n,;।]+",
        r"(?:সঠিক\s+উত্তর|উত্তর)\s*[:：\-]?\s*[^\n,;।]+",
        r"(?:উ)\s*[\.:：\-]\s*[কখগঘabcdABCD]",
    ]

    return any(re.search(pattern, text_value, flags=re.IGNORECASE) for pattern in markers)


def normalize_answer_to_option_number(answer, options):
    if answer is None:
        return None

    if isinstance(answer, int) and 0 <= answer <= 4:
        return answer

    answer_text = str(answer).strip()
    if not answer_text:
        return None

    if answer_text.isdigit():
        answer_number = int(answer_text)
        if 0 <= answer_number <= 4:
            return answer_number

    label_map = {
        "a": 1,
        "b": 2,
        "c": 3,
        "d": 4,
        "ক": 1,
        "খ": 2,
        "গ": 3,
        "ঘ": 4,
    }
    cleaned_label = answer_text.strip(" .):।-").lower()
    cleaned_label = re.sub(r"^(?:উ|উত্তর|ans|answer)\s*[\.:：\-]?\s*", "", cleaned_label, flags=re.IGNORECASE)
    if cleaned_label in label_map:
        return label_map[cleaned_label]

    normalized_answer = re.sub(r"\s+", " ", answer_text).strip().lower()
    for index, option in enumerate(options, start=1):
        option_text = ""
        if isinstance(option, dict):
            option_text = option.get("text") or ""
        else:
            option_text = str(option or "")

        normalized_option = re.sub(r"\s+", " ", str(option_text)).strip().lower()
        if normalized_option and normalized_option == normalized_answer:
            return index

    return None


def normalize_mcq_dataset(value, image_lookup=None):
    if not isinstance(value, dict):
        raise ValueError("Structured extraction did not return a JSON object.")

    raw_mcqs = value.get("mcqs", [])
    raw_contexts = value.get("contexts", [])
    if not isinstance(raw_mcqs, list):
        raw_mcqs = []
    if not isinstance(raw_contexts, list):
        raw_contexts = []

    normalized_mcqs = []
    for item in raw_mcqs:
        if not isinstance(item, dict):
            continue

        question = str(item.get("question") or "").strip()
        options = item.get("options")
        if not question or not isinstance(options, list) or len(options) != 4:
            continue

        normalized_options = []
        for option in options:
            if isinstance(option, dict):
                option_type = str(option.get("type") or "text").strip().lower()
                text = option.get("text")
                text = str(text).strip() if text is not None and str(text).strip() else None
                image_base64 = option.get("imageBase64") or option.get("image_base64") or option.get("image")
                image_base64 = str(image_base64).strip() if image_base64 else None
                image_ref = option.get("imageRef") or option.get("image_ref")
                image_ref = str(image_ref).strip() if image_ref else None

                if image_lookup:
                    if image_base64 and not image_base64.startswith("data:image/"):
                        image_base64 = image_lookup.get(image_base64) or image_base64
                    if not image_base64 and image_ref:
                        image_base64 = image_lookup.get(image_ref)

                if option_type not in {"text", "image"}:
                    option_type = "image" if image_base64 else "text"

                if option_type == "image" and not image_base64:
                    continue
                if option_type == "text" and not text:
                    continue

                normalized_options.append({
                    "type": option_type,
                    "text": text,
                    "imageBase64": image_base64,
                })
            else:
                text = str(option or "").strip()
                if text:
                    normalized_options.append({
                        "type": "text",
                        "text": text,
                        "imageBase64": None,
                    })

        if len(normalized_options) != 4:
            continue

        explanation = item.get("explanation")
        if isinstance(explanation, str):
            explanation = explanation.strip() or None
        elif explanation is not None:
            explanation = str(explanation).strip() or None

        raw_answer = item.get("answer")
        answer = normalize_answer_to_option_number(raw_answer, normalized_options)
        if answer is None and explanation:
            answer = extract_option_number_from_text(explanation, normalized_options)
            if answer is None and contains_answer_evidence(explanation):
                answer = 0
        elif answer is None and contains_answer_evidence(raw_answer):
            answer = 0

        language = item.get("language")
        if language not in {"bn", "en", "mixed"}:
            option_text = " ".join(option.get("text") or "" for option in normalized_options)
            language = detect_text_language(" ".join([question, option_text]))

        subject = item.get("subject")
        subject = str(subject).strip() if subject is not None and str(subject).strip() else None

        exam = item.get("exam")
        exam = str(exam).strip() if exam is not None and str(exam).strip() else None

        year = item.get("year")
        year = str(year).strip() if year is not None and str(year).strip() else None

        normalized_mcqs.append({
            "question": question,
            "options": normalized_options,
            "answer": answer,
            "explanation": explanation,
            "language": language,
            "subject": subject,
            "exam": exam,
            "year": year,
        })

    normalized_contexts = [
        str(context).strip()
        for context in raw_contexts
        if isinstance(context, str) and context.strip()
    ]

    return {
        "mcqs": normalized_mcqs,
        "contexts": normalized_contexts,
    }


def detect_text_language(text):
    has_bangla = bool(re.search(r"[\u0980-\u09FF]", text or ""))
    has_latin = bool(re.search(r"[A-Za-z]", text or ""))
    if has_bangla and has_latin:
        return "mixed"
    if has_bangla:
        return "bn"
    return "en"


def transform_ocr_text_to_mcq_dataset(raw_text, model):
    if not raw_text or not raw_text.strip():
        return {"mcqs": [], "contexts": []}

    response = call_openai_compatible_chat(
        model,
        messages=[
            {"role": "system", "content": MCQ_DATASET_SYSTEM_PROMPT},
            {
                "role": "user",
                "content": (
                    "Convert the following OCR text into the required strict JSON dataset.\n\n"
                    f"{raw_text}"
                ),
            },
        ],
        temperature=0,
    )

    message = get_openai_compatible_message_text(response)

    try:
        parsed = parse_json_llm_response(message)
    except json.JSONDecodeError as exc:
        raise ValueError("Structured MCQ extraction returned invalid JSON.") from exc

    return normalize_mcq_dataset(parsed)


def get_image_mime_type(path_or_url):
    ext = os.path.splitext((path_or_url or "").split("?", 1)[0])[1].lower()
    if ext in {".jpg", ".jpeg"}:
        return "image/jpeg"
    if ext == ".png":
        return "image/png"
    if ext == ".webp":
        return "image/webp"
    if ext in {".tif", ".tiff"}:
        return "image/tiff"
    if ext == ".bmp":
        return "image/bmp"
    return "image/jpeg"


def image_bytes_to_data_url(image_bytes, path_or_url):
    encoded = base64.b64encode(image_bytes).decode("ascii")
    return f"data:{get_image_mime_type(path_or_url)};base64,{encoded}"


def download_image_bytes(url):
    response = requests.get(url, timeout=60)
    response.raise_for_status()
    return response.content


def replace_markdown_image_sources_with_refs(markdown_text, page_number, page_images):
    text = markdown_text or ""
    for img_path in page_images.keys():
        image_ref = f"page_{page_number}:{img_path}"
        text = text.replace(f'src="{img_path}"', f'src="{image_ref}"')
        text = text.replace(f"src='{img_path}'", f"src='{image_ref}'")
    return text


def build_source_page_image_assets(file_path):
    assets = []

    if is_pdf_file(file_path):
        doc = fitz.open(file_path)
        try:
            for page_index, page in enumerate(doc, start=1):
                pix = page.get_pixmap(dpi=180)
                image = Image.open(io.BytesIO(pix.tobytes("png")))
                assets.append({
                    "id": f"page_{page_index}:full_page",
                    "path": f"page_{page_index}.jpg",
                    "data_url": prepare_openai_compatible_image_data_url(image),
                    "kind": "full_page",
                })
        finally:
            doc.close()
    elif is_image_file(file_path):
        image = Image.open(file_path)
        assets.append({
            "id": "page_1:full_page",
            "path": os.path.basename(file_path),
            "data_url": prepare_openai_compatible_image_data_url(image),
            "kind": "full_page",
        })

    return assets


def transform_aistudio_layout_to_mcq_dataset(markdown_text, image_assets, model, max_images=80):
    if not markdown_text or not markdown_text.strip():
        return {"mcqs": [], "contexts": []}

    image_lookup = {
        asset["id"]: asset["data_url"]
        for asset in image_assets
        if asset.get("id") and asset.get("data_url")
    }

    content = [
        {
            "type": "text",
            "text": (
                "Convert this AI Studio/PaddleOCR layout output into the required strict JSON dataset.\n\n"
                "The markdown contains image placeholders like src=\"page_1:imgs/example.jpg\". "
                "Those image IDs are attached after the markdown as image inputs. "
                "Full-page images are also attached with IDs like page_1:full_page. "
                "Use the full-page images to read right-margin printed answer markers such as "
                "\"উ. ক\", \"উ. খ\", \"উ. গ\", or \"উ. ঘ\"; map them to answer 1, 2, 3, or 4. "
                "Use both OCR text and attached images. If an MCQ option is visual, set that option to "
                "{\"type\":\"image\",\"text\":<label or null>,\"imageRef\":\"<matching image id>\"}. "
                "If an option has text only, use {\"type\":\"text\",\"text\":\"...\",\"imageBase64\":null}. "
                "Do not leave answer null when a printed answer marker is visible in the full-page image. "
                "Return only JSON; do not include markdown or explanation.\n\n"
                f"{markdown_text}"
            ),
        }
    ]

    ordered_assets = sorted(
        image_assets,
        key=lambda asset: 0 if asset.get("kind") == "full_page" else 1
    )

    for asset in ordered_assets[:max_images]:
        data_url = asset.get("data_url")
        image_id = asset.get("id")
        if not data_url or not image_id:
            continue
        content.append({"type": "text", "text": f"Image ID: {image_id}"})
        content.append({"type": "image_url", "image_url": {"url": data_url}})

    response = call_openai_compatible_chat(
        model,
        messages=[
            {"role": "system", "content": MCQ_DATASET_SYSTEM_PROMPT},
            {"role": "user", "content": content},
        ],
        temperature=0,
    )

    message = get_openai_compatible_message_text(response)

    try:
        parsed = parse_json_llm_response(message)
    except json.JSONDecodeError as exc:
        raise ValueError("Structured MCQ extraction returned invalid JSON.") from exc

    return normalize_mcq_dataset(parsed, image_lookup=image_lookup)


def repair_aistudio_mcq_with_page_image(mcq, page_text, page_assets, model, max_images=80):
    if not mcq or not page_assets:
        return mcq

    image_lookup = {
        asset["id"]: asset["data_url"]
        for asset in page_assets
        if asset.get("id") and asset.get("data_url")
    }

    content = [
        {
            "type": "text",
            "text": (
                "Repair this single MCQ extraction by re-reading the attached page image and layout text. "
                "The previous extraction has a missing or contradictory answer. "
                "If the page shows a right-margin marker like \"উ. ক\", \"উ. খ\", \"উ. গ\", or \"উ. ঘ\", "
                "map it to answer 1, 2, 3, or 4. This right-margin marker is authoritative. "
                "Ignore the previous options if they conflict with the page image. If an explanation says "
                "\"সঠিক উত্তর: <value>\", make sure the option containing that value is captured from the page image "
                "and use its option number. If the explanation answer is visible but not present in any option, "
                "set answer to 0. Correct OCR mistakes in the option text from the page image. "
                "Return ONLY valid JSON with this exact shape: "
                "{\"mcqs\":[{\"question\":\"\",\"options\":[{\"type\":\"text\",\"text\":\"\",\"imageBase64\":null}],"
                "\"answer\":1,\"explanation\":null,\"language\":\"bn\",\"subject\":null,\"exam\":null}],\"contexts\":[]}.\n\n"
                "Previous MCQ JSON:\n"
                f"{json.dumps(mcq, ensure_ascii=False)}\n\n"
                "Page layout text:\n"
                f"{page_text}"
            ),
        }
    ]

    ordered_assets = sorted(
        page_assets,
        key=lambda asset: 0 if asset.get("kind") == "full_page" else 1
    )
    for asset in ordered_assets[:max_images]:
        data_url = asset.get("data_url")
        image_id = asset.get("id")
        if not data_url or not image_id:
            continue
        content.append({"type": "text", "text": f"Image ID: {image_id}"})
        content.append({"type": "image_url", "image_url": {"url": data_url}})

    response = call_openai_compatible_chat(
        model,
        messages=[
            {"role": "system", "content": MCQ_DATASET_SYSTEM_PROMPT},
            {"role": "user", "content": content},
        ],
        temperature=0,
    )

    message = get_openai_compatible_message_text(response)
    try:
        parsed = parse_json_llm_response(message)
    except json.JSONDecodeError:
        return mcq

    repaired = normalize_mcq_dataset(parsed, image_lookup=image_lookup).get("mcqs", [])
    if repaired:
        return repaired[0]

    return mcq


def merge_mcq_datasets(datasets):
    merged_mcqs = []
    merged_contexts = []
    seen_questions = set()

    for dataset in datasets:
        for mcq in dataset.get("mcqs", []):
            question_key = re.sub(r"\s+", " ", mcq.get("question") or "").strip()
            if not question_key or question_key in seen_questions:
                continue
            seen_questions.add(question_key)
            merged_mcqs.append(mcq)

        for context in dataset.get("contexts", []):
            if context and context not in merged_contexts:
                merged_contexts.append(context)

    numbered_count = sum(
        1 for mcq in merged_mcqs
        if get_question_number_from_text(mcq.get("question") or "") is not None
    )
    if numbered_count >= max(2, len(merged_mcqs) // 2):
        merged_mcqs.sort(
            key=lambda mcq: (
                get_question_number_from_text(mcq.get("question") or "") is None,
                get_question_number_from_text(mcq.get("question") or "") or 10**9,
            )
        )

    return {
        "mcqs": merged_mcqs,
        "contexts": merged_contexts,
    }



# ==========================================
# Anthropic Claude Compatibility Layer
# ==========================================

def call_anthropic_api(model, messages, temperature=0, timeout=None):
    import anthropic
    config = get_config()
    api_key = config.get('ANTHROPIC_API_KEY', '')
    base_url = config.get('ANTHROPIC_BASE_URL', '')
    if not api_key:
        raise ValueError("Anthropic API key not configured. Set ANTHROPIC_API_KEY in settings.")

    client_kwargs = {
        "api_key": api_key,
        "timeout": timeout or float(config.get('ANTHROPIC_TIMEOUT_SECONDS', 240))
    }
    if base_url:
        client_kwargs["base_url"] = base_url

    client = anthropic.Anthropic(**client_kwargs)

    system_prompt = ""
    anthropic_messages = []

    for msg in messages:
        if msg['role'] == 'system':
            system_prompt += msg['content'] + "\n"
        elif msg['role'] in ['user', 'assistant']:
            if isinstance(msg['content'], list):
                new_content = []
                for item in msg['content']:
                    if item['type'] == 'text':
                        new_content.append({"type": "text", "text": item['text']})
                    elif item['type'] == 'image_url':
                        url = item['image_url']['url']
                        media_type, base64_data = url.split(";base64,")
                        media_type = media_type.replace("data:", "")
                        new_content.append({
                            "type": "image",
                            "source": {
                                "type": "base64",
                                "media_type": media_type,
                                "data": base64_data
                            }
                        })
                anthropic_messages.append({"role": msg['role'], "content": new_content})
            else:
                anthropic_messages.append({"role": msg['role'], "content": msg['content']})

    response = client.messages.create(
        model=model or "claude-3-5-sonnet-latest",
        system=system_prompt.strip() if system_prompt else anthropic.NOT_GIVEN,
        messages=anthropic_messages,
        max_tokens=4096,
        temperature=temperature
    )
    
    text_parts = []
    for block in response.content:
        if hasattr(block, 'text') and block.text:
            text_parts.append(block.text)
        elif hasattr(block, 'thinking') and block.thinking:
            text_parts.append(block.thinking)
            
    result = "".join(text_parts).strip()
    if not result:
        # Prevent json.loads expecting value error by throwing a clear runtime error
        raw_content = getattr(response, 'content', 'No content')
        raise RuntimeError(f"Anthropic API returned an empty text response. Raw blocks: {raw_content}")
    return result

def extract_clean_text_with_anthropic(image, page_number, model):
    data_url = prepare_openai_compatible_image_data_url(image)

    message = call_anthropic_api(
        model,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            "OCR this document page and clean the result. Preserve Bangla and English exactly. "
                            "Keep question numbers, options, formulas, answer markers such as 'উ. ক', and explanations. "
                            "Do not summarize. Return only cleaned page text."
                        ),
                    },
                    {"type": "image_url", "image_url": {"url": data_url}},
                ],
            }
        ],
        temperature=0,
    )

    if message:
        return message
    raise RuntimeError(f"No text returned by Anthropic engine for page {page_number}.")

def transform_page_image_to_mcq_dataset_with_anthropic(image, page_number, model, page_text=None):
    data_url = prepare_openai_compatible_image_data_url(image)

    message = call_anthropic_api(
        model,
        messages=[
            {"role": "system", "content": MCQ_DATASET_SYSTEM_PROMPT},
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            f"OCR, clean, and structure all complete MCQs visible on page {page_number}. "
                            "Use both the cleaned OCR text below and the page image directly for OCR, options, formulas, "
                            "answer markers, and explanations. "
                            "Do not omit visible numbered questions. Right-side printed markers like 'উ. ক', 'উ. খ', "
                            "'উ: গ', 'উত্তর: ঘ' usually denote the correct answer option.\n\n"
                            f"=== Cleaned OCR Text (Page {page_number}) ===\n{page_text or 'No text provided.'}"
                        ),
                    },
                    {"type": "image_url", "image_url": {"url": data_url}},
                ],
            }
        ],
        temperature=0,
    )

    return normalize_mcq_dataset(parse_json_llm_response(message))

def extract_missing_mcqs_with_anthropic(image, page_number, page_text, existing_dataset, model):
    data_url = prepare_openai_compatible_image_data_url(image)
    existing_mcqs_json = json.dumps(existing_dataset.get("mcqs", []), ensure_ascii=False, indent=2)

    message = call_anthropic_api(
        model,
        messages=[
            {"role": "system", "content": MCQ_DATASET_SYSTEM_PROMPT},
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            f"You previously extracted MCQs from page {page_number}, but some are missing. "
                            "Look at the page image and the extracted text again. "
                            f"These are the MCQs you already found:\n```json\n{existing_mcqs_json}\n```\n\n"
                            "Extract ONLY the MCQs that are present in the image/text but MISSING from the JSON above. "
                            "Do not repeat already extracted MCQs. "
                            f"=== Cleaned OCR Text (Page {page_number}) ===\n{page_text or 'No text provided.'}"
                        ),
                    },
                    {"type": "image_url", "image_url": {"url": data_url}},
                ],
            }
        ],
        temperature=0,
    )

    return normalize_mcq_dataset(parse_json_llm_response(message))

def process_with_anthropic(file_path, output_format, job_id=None, page_range=None):
    config = get_config()
    def update_progress(percentage, status):
        if job_id:
            job_progress[job_id] = {"status": status, "percentage": percentage}

    def check_controls():
        if job_id and job_id in job_controls:
            if job_controls[job_id].get('cancel_flag'):
                raise InterruptedError("Job cancelled by user")
            job_controls[job_id]['pause_event'].wait()

    update_progress(5, "Initializing Anthropic engine...")
    model = config.get('ANTHROPIC_MODEL', 'claude-3-5-sonnet-latest')

    total_pdf_pages = 0
    if is_pdf_file(file_path):
        doc = fitz.open(file_path)
        total_pdf_pages = len(doc)
        doc.close()

    page_indices = parse_page_range(page_range, total_pdf_pages) if is_pdf_file(file_path) else None
    page_images = get_document_page_images(file_path, dpi=int(config.get('OCR_DPI', 300)), page_indices=page_indices)
    total_pages = len(page_images)

    if output_format == "json":
        page_datasets = []
        for idx, (page_number, image) in enumerate(page_images):
            check_controls()
            update_progress(
                5 + int((idx / total_pages) * 90),
                f"Processing page {page_number} of {total_pages}..."
            )
            page_text = extract_clean_text_with_anthropic(image, page_number, model)
            page_dataset = transform_page_image_to_mcq_dataset_with_anthropic(
                image, page_number, model, page_text=page_text
            )
            expected_count = get_longest_consecutive_question_count(page_text)
            if expected_count and len(page_dataset.get("mcqs", [])) < expected_count:
                update_progress(
                    5 + int((idx / total_pages) * 90) + 2,
                    f"Repairing page {page_number} (found {len(page_dataset.get('mcqs', []))}/{expected_count})..."
                )
                missing_dataset = extract_missing_mcqs_with_anthropic(
                    image, page_number, page_text, page_dataset, model
                )
                page_dataset = merge_mcq_datasets([page_dataset, missing_dataset])
            page_datasets.append(page_dataset)

        dataset = merge_mcq_datasets(page_datasets)
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], f'anthropic_{job_id or uuid.uuid4().hex}.json')
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(dataset, f, ensure_ascii=False, indent=2)
        preview_text = json.dumps(dataset, ensure_ascii=False, indent=2)
    else:
        extracted_text = []
        for idx, (page_number, image) in enumerate(page_images):
            update_progress(
                5 + int((idx / total_pages) * 90),
                f"Processing page {page_number} of {total_pages}..."
            )
            text = extract_clean_text_with_anthropic(image, page_number, model)
            extracted_text.append({
                "page": page_number,
                "text": text or "",
            })

        full_text = "\n\n".join(page["text"] for page in extracted_text)

        if output_format == "txt":
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], f'anthropic_{job_id or uuid.uuid4().hex}.txt')
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(full_text)
            preview_text = full_text
        elif output_format == "docx":
            from docx import Document
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], f'anthropic_{job_id or uuid.uuid4().hex}.docx')
            doc = Document()
            doc.add_heading('Anthropic Extracted Text', 0)
            for page_data in extracted_text:
                doc.add_heading(f'Page {page_data["page"]}', level=1)
                for line_text in page_data["text"].splitlines():
                    doc.add_paragraph(line_text)
            doc.save(output_path)
            preview_text = full_text
        else:
            raise ValueError("Anthropic engine supports TXT, DOCX, and JSON output formats.")

    update_progress(100, "Done!")
    return {
        'text': preview_text,
        'pages': total_pages,
        'output_path': output_path
    }


def process_with_gemini(file_path, output_format, job_id=None, page_range=None):
    config = get_config()
    def update_progress(percentage, status):
        if job_id:
            job_progress[job_id] = {"status": status, "percentage": percentage}

    def check_controls():
        if job_id and job_id in job_controls:
            if job_controls[job_id].get('cancel_flag'):
                raise InterruptedError("Job cancelled by user")
            job_controls[job_id]['pause_event'].wait()

    update_progress(5, "Initializing Gemini engine...")
    caps = get_openai_compat_capabilities()
    if not caps["vision_supported"]:
        raise RuntimeError(
            "The configured local OpenAI-compatible model does not support images. "
            "Set OPENAI_COMPAT_MODEL to a Gemini vision-capable model."
        )

    model = caps["vision_model"]

    # Handle page range
    total_pdf_pages = 0
    if is_pdf_file(file_path):
        doc = fitz.open(file_path)
        total_pdf_pages = len(doc)
        doc.close()

    page_indices = parse_page_range(page_range, total_pdf_pages) if is_pdf_file(file_path) else None
    page_images = get_document_page_images(file_path, dpi=config['OCR_DPI'], page_indices=page_indices)
    total_pages = len(page_images)

    if output_format == "json":
        page_datasets = []
        for idx, (page_number, image) in enumerate(page_images):
            check_controls()
            update_progress(
                5 + int((idx / total_pages) * 90),
                f"Processing page {page_number} of {total_pages}..."
            )
            page_text = extract_clean_text_with_gemini(image, page_number, model)
            page_dataset = transform_page_image_to_mcq_dataset_with_gemini(
                image,
                page_number,
                model,
                page_text=page_text,
            )
            expected_count = get_longest_consecutive_question_count(page_text)
            if expected_count and len(page_dataset.get("mcqs", [])) < expected_count:
                update_progress(
                    5 + int((idx / total_pages) * 90) + 2,
                    f"Repairing page {page_number} (found {len(page_dataset.get('mcqs', []))}/{expected_count})..."
                )
                missing_dataset = extract_missing_mcqs_with_gemini(
                    image,
                    page_number,
                    page_text,
                    page_dataset,
                    model,
                )
                page_dataset = merge_mcq_datasets([page_dataset, missing_dataset])
            page_datasets.append(page_dataset)
        dataset = merge_mcq_datasets(page_datasets)
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], f'gemini_{job_id or uuid.uuid4().hex}.json')
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(dataset, f, ensure_ascii=False, indent=2)
        preview_text = json.dumps(dataset, ensure_ascii=False, indent=2)
    else:
        extracted_text = []
        for idx, (page_number, image) in enumerate(page_images):
            update_progress(
                5 + int((idx / total_pages) * 90),
                f"Processing page {page_number} of {total_pages}..."
            )
            text = extract_clean_text_with_gemini(image, page_number, model)
            extracted_text.append({
                "page": page_number,
                "text": text or "",
            })

        full_text = "\n\n".join(page["text"] for page in extracted_text)

        if output_format == "txt":
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], f'gemini_{job_id or uuid.uuid4().hex}.txt')
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(full_text)
            preview_text = full_text
        elif output_format == "docx":
            from docx import Document
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], f'gemini_{job_id or uuid.uuid4().hex}.docx')
            doc = Document()
            doc.add_heading('Gemini Extracted Text', 0)
            for page_data in extracted_text:
                doc.add_heading(f'Page {page_data["page"]}', level=1)
                for line_text in page_data["text"].splitlines():
                    doc.add_paragraph(line_text)
            doc.save(output_path)
            preview_text = full_text
        else:
            raise ValueError("Gemini engine supports TXT, DOCX, and JSON output formats.")

    update_progress(100, "Done!")

    return {
        'text': preview_text,
        'pages': len(page_images),
        'output_path': output_path,
    }

def configure_tesseract_command(pytesseract_module):
    config = get_config()
    custom_path = config.get('TESSERACT_PATH')
    if custom_path and os.path.exists(custom_path):
        pytesseract_module.pytesseract.tesseract_cmd = custom_path
        return custom_path

    # Use PATH first; if not found, try common Windows install locations.
    tesseract_on_path = shutil.which("tesseract")
    if tesseract_on_path:
        pytesseract_module.pytesseract.tesseract_cmd = tesseract_on_path
        return tesseract_on_path

    if os.name == "nt":
        candidates = [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        ]
        for candidate in candidates:
            if os.path.exists(candidate):
                pytesseract_module.pytesseract.tesseract_cmd = candidate
                return candidate

    return None


def html_block_to_plain_text(content):
    if not content:
        return ""
    text = content
    text = re.sub(r'(?i)<br\\s*/?>', '\n', text)
    text = re.sub(r'(?i)</tr\\s*>', '\n', text)
    text = re.sub(r'(?i)</t[dh]\\s*>', '\t', text)
    text = re.sub(r'<[^>]+>', '', text)
    text = unescape(text)
    lines = [line.strip() for line in text.splitlines()]
    return "\n".join([line for line in lines if line])


def create_positioned_docx_from_aistudio(layout_results, output_path):
    from docx import Document
    from docx.shared import Pt, Twips
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Keep page width standard Letter; scale all coordinates proportionally.
    target_page_width_twips = 12240
    max_page_height_twips = 30000

    doc = Document()
    first_page = True

    for page_result in layout_results:
        pruned = page_result.get('prunedResult', {}) or {}
        page_width = max(1, int(pruned.get('width') or 1))
        page_height = max(1, int(pruned.get('height') or 1))
        blocks = pruned.get('parsing_res_list') or []
        blocks = sorted(
            blocks,
            key=lambda b: b.get('block_order') if isinstance(b.get('block_order'), int) else 10**9
        )

        scale = target_page_width_twips / page_width
        page_height_twips = int(page_height * scale)
        if page_height_twips > max_page_height_twips:
            scale = max_page_height_twips / page_height
            page_height_twips = max_page_height_twips

        if first_page:
            section = doc.sections[0]
            first_page = False
        else:
            section = doc.add_section(start_type=1)

        section.page_width = Twips(target_page_width_twips)
        section.page_height = Twips(max(1, page_height_twips))
        section.left_margin = Twips(0)
        section.right_margin = Twips(0)
        section.top_margin = Twips(0)
        section.bottom_margin = Twips(0)

        for block in blocks:
            bbox = block.get('block_bbox')
            text = html_block_to_plain_text(block.get('block_content', ''))
            if not text or not isinstance(bbox, list) or len(bbox) != 4:
                continue

            x1, y1, x2, y2 = bbox
            w = max(1, int((x2 - x1) * scale))
            h = max(1, int((y2 - y1) * scale))
            x = max(0, int(x1 * scale))
            y = max(0, int(y1 * scale))

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0

            p_pr = p._p.get_or_add_pPr()
            frame_pr = OxmlElement('w:framePr')
            frame_pr.set(qn('w:w'), str(w))
            frame_pr.set(qn('w:h'), str(h))
            frame_pr.set(qn('w:x'), str(x))
            frame_pr.set(qn('w:y'), str(y))
            frame_pr.set(qn('w:hRule'), 'exact')
            frame_pr.set(qn('w:vAnchor'), 'page')
            frame_pr.set(qn('w:hAnchor'), 'page')
            frame_pr.set(qn('w:wrap'), 'none')
            p_pr.insert(0, frame_pr)

            run = p.add_run(text)
            run.font.name = 'Calibri'
            run.font.size = Pt(10)

    doc.save(output_path)


def process_with_aistudio(file_path, output_format, job_id=None, page_range=None):
    config = get_config()
    API_URL = config['AISTUDIO_API_URL']
    TOKEN = config['AISTUDIO_TOKEN']

    def update_progress(percentage, status):
        if job_id:
            job_progress[job_id] = {"status": status, "percentage": percentage}

    def check_controls():
        if job_id and job_id in job_controls:
            if job_controls[job_id].get('cancel_flag'):
                raise InterruptedError("Job cancelled by user")
            job_controls[job_id]['pause_event'].wait()

    update_progress(5, "Initializing AI Studio...")

    # We will split the PDF into individual pages to avoid timeouts on large files
    # and to provide real-time progress updates.
    all_layout_results = []

    try:
        page_contents = [] # List of (page_number, base64_data, is_pdf)
        if is_pdf_file(file_path):
            doc = fitz.open(file_path)
            total_pdf_pages = len(doc)

            page_indices = parse_page_range(page_range, total_pdf_pages)
            for i in [idx - 1 for idx in page_indices]: # Convert 1-based to 0-based
                new_doc = fitz.open()
                new_doc.insert_pdf(doc, from_page=i, to_page=i)
                # Save small PDF to memory
                pdf_bytes = new_doc.tobytes()
                new_doc.close()

                file_data = base64.b64encode(pdf_bytes).decode("ascii")
                page_contents.append((i + 1, file_data, True))
            doc.close()
        else:
            with open(file_path, "rb") as file:
                file_bytes = file.read()
                file_data = base64.b64encode(file_bytes).decode("ascii")
            page_contents.append((1, file_data, False))

        total_pages = len(page_contents)

        for idx, (page_num, file_data, is_pdf_page) in enumerate(page_contents):
            check_controls()
            update_progress(
                5 + int((idx / total_pages) * 80),
                f"Processing page {page_num} of {total_pages}..."
            )

            headers = {
                "Authorization": f"token {TOKEN}",
                "Content-Type": "application/json"
            }

            payload = {
                "file": file_data,
                "fileType": 0 if is_pdf_page else 1,
                "useDocOrientationClassify": False,
                "useDocUnwarping": False,
                "useChartRecognition": False,
            }

            # Use a smaller timeout per page, but retry if needed
            max_retries = config['OCR_MAX_RETRIES']
            for attempt in range(max_retries):
                try:
                    response = requests.post(API_URL, json=payload, headers=headers, timeout=config['AISTUDIO_TIMEOUT_SECONDS'])
                    if response.status_code == 200:
                        page_result = response.json().get("result", {})
                        all_layout_results.extend(page_result.get("layoutParsingResults", []))
                        break
                    else:
                        if attempt == max_retries - 1:
                            raise Exception(f"API Error {response.status_code}: {response.text}")
                except requests.Timeout:
                    if attempt == max_retries - 1:
                        raise Exception(f"AI Studio request timed out on page {page_num}")
                except Exception as e:
                    if attempt == max_retries - 1:
                        raise e

        update_progress(85, "Merging results and generating output...")

        # Now we have all layout results, we can proceed with the rest of the original logic
        # but using all_layout_results instead of result["layoutParsingResults"]

        output_dir = os.path.join(app.config['OUTPUT_FOLDER'], f"aistudio_{job_id or uuid.uuid4().hex}")
        os.makedirs(output_dir, exist_ok=True)

        full_markdown_text = ""
        image_assets = []
        page_payloads = []
        full_page_assets = build_source_page_image_assets(file_path)
        pages = 0

        for i, res in enumerate(all_layout_results):
            pages += 1
            page_assets = [
                asset for asset in full_page_assets
                if asset.get("id") == f"page_{pages}:full_page"
            ]
            md_filename = os.path.join(output_dir, f"doc_{i}.md")
            markdown = res.get("markdown", {}) or {}
            page_images = markdown.get("images", {}) or {}
            page_md = markdown.get("text", "")
            page_md_for_gemini = replace_markdown_image_sources_with_refs(page_md, pages, page_images)
            full_markdown_text += page_md_for_gemini + "\n\n"
            with open(md_filename, "w", encoding="utf-8") as md_file:
                md_file.write(page_md)

            for img_path, img_url in page_images.items():
                full_img_path = os.path.join(output_dir, img_path)
                os.makedirs(os.path.dirname(full_img_path), exist_ok=True)
                try:
                    img_bytes = download_image_bytes(img_url)
                    with open(full_img_path, "wb") as img_file:
                        img_file.write(img_bytes)
                    image_assets.append({
                        "id": f"page_{pages}:{img_path}",
                        "path": img_path,
                        "data_url": image_bytes_to_data_url(img_bytes, img_path),
                    })
                    page_assets.append(image_assets[-1])
                except Exception:
                    pass

            if "outputImages" in res:
                for img_name, img_url in res["outputImages"].items():
                    try:
                        img_response = requests.get(img_url, timeout=60)
                        if img_response.status_code == 200:
                            filename = os.path.join(output_dir, f"{img_name}_{i}.jpg")
                            with open(filename, "wb") as f:
                                f.write(img_response.content)
                    except:
                        pass

            page_payloads.append({
                "text": page_md_for_gemini,
                "assets": page_assets,
            })

        # Create zip archive of the directory
        zip_path_base = os.path.join(app.config['OUTPUT_FOLDER'], f"aistudio_{job_id or uuid.uuid4().hex}")
        shutil.make_archive(zip_path_base, 'zip', output_dir)

        if output_format == 'json':
            update_progress(90, "Converting to JSON dataset...")
            caps = get_openai_compat_capabilities()
            model = caps["vision_model"] or caps["text_model"]
            if caps["vision_supported"]:
                page_datasets = []
                for idx, page_payload in enumerate(page_payloads):
                    update_progress(90 + int((idx/len(page_payloads))*9), f"Analyzing page {idx+1} layout...")
                    page_text = page_payload.get("text") or ""
                    page_assets_for_model = page_payload.get("assets") or []
                    if not page_text.strip():
                        continue
                    page_dataset = transform_aistudio_layout_to_mcq_dataset(page_text, page_assets_for_model, model, max_images=config['AISTUDIO_MAX_GEMINI_IMAGES'])
                    repaired_mcqs = []
                    for mcq in page_dataset.get("mcqs", []):
                        if mcq.get("answer") is None:
                            mcq = repair_aistudio_mcq_with_page_image(
                                mcq,
                                page_text,
                                page_assets_for_model,
                                model,
                                max_images=config['AISTUDIO_MAX_GEMINI_IMAGES']
                            )
                        repaired_mcqs.append(mcq)
                    page_dataset["mcqs"] = repaired_mcqs
                    page_datasets.append(page_dataset)
                dataset = merge_mcq_datasets(page_datasets)
            else:
                dataset = transform_ocr_text_to_mcq_dataset(full_markdown_text, model)
            final_output_path = os.path.join(app.config['OUTPUT_FOLDER'], f"aistudio_{job_id or uuid.uuid4().hex}.json")
            with open(final_output_path, 'w', encoding='utf-8') as f:
                json.dump(dataset, f, ensure_ascii=False, indent=2)
        elif output_format == 'zip':
            final_output_path = zip_path_base + ".zip"
        elif output_format == 'txt':
            final_output_path = os.path.join(app.config['OUTPUT_FOLDER'], f"aistudio_{job_id or uuid.uuid4().hex}.txt")
            with open(final_output_path, 'w', encoding='utf-8') as f:
                f.write(full_markdown_text)
        elif output_format == 'docx':
            final_output_path = os.path.join(app.config['OUTPUT_FOLDER'], f"aistudio_{job_id or uuid.uuid4().hex}.docx")
            create_positioned_docx_from_aistudio(all_layout_results, final_output_path)
        else:
            final_output_path = zip_path_base + ".zip"

        update_progress(100, "Done!")

        # Optional cleanup
        try:
            shutil.rmtree(output_dir)
        except:
            pass

        return {
            'text': json.dumps(dataset, ensure_ascii=False, indent=2) if output_format == 'json' else full_markdown_text,
            'pages': pages,
            'output_path': final_output_path
        }
    except Exception as e:
        raise e


def reconstruct_layout(ocr_results):
    if not ocr_results:
        return ""
        
    boxes = []
    for res in ocr_results:
        if len(res) == 3:
            bounds, text, conf = res
            if not bounds or not isinstance(bounds, list) or len(bounds) != 4:
                continue
            
            try:
                min_x = min([p[0] for p in bounds])
                max_x = max([p[0] for p in bounds])
                min_y = min([p[1] for p in bounds])
                max_y = max([p[1] for p in bounds])
                
                # Approximate width per character
                char_width = (max_x - min_x) / max(1, len(text))
                
                boxes.append({
                    'min_x': min_x, 'max_x': max_x, 
                    'min_y': min_y, 'max_y': max_y,
                    'text': text, 'char_width': char_width
                })
            except Exception:
                pass

    if not boxes:
        return ""

    # Sort boxes by top y coordinate
    boxes.sort(key=lambda b: b['min_y'])

    lines = []
    current_line = []
    
    for box in boxes:
        if not current_line:
            current_line.append(box)
        else:
            # Check if this box belongs to the current line based on Y-overlap
            min_y_line = min([b['min_y'] for b in current_line])
            max_y_line = max([b['max_y'] for b in current_line])
            line_height = max(1, max_y_line - min_y_line)
            
            box_center_y = (box['min_y'] + box['max_y']) / 2
            
            # If the box center is within the current line vertical bounds, or very close
            if min_y_line <= box_center_y <= max_y_line or abs(box_center_y - (min_y_line + max_y_line)/2) < line_height * 0.5:
                current_line.append(box)
            else:
                lines.append(current_line)
                current_line = [box]
                
    if current_line:
        lines.append(current_line)
        
    # Build text preserving approximate horizontal gaps
    output = []
    for line in lines:
        line.sort(key=lambda b: b['min_x'])
        
        char_widths = [b['char_width'] for b in line if b['char_width'] > 0]
        avg_char_width = sum(char_widths) / len(char_widths) if char_widths else 10
        avg_char_width = max(1, avg_char_width)
        
        line_str = ""
        last_x = 0
        
        for i, box in enumerate(line):
            gap = box['min_x'] - last_x
            
            if i == 0:
                # Add initial indent relative to left margin
                if gap > avg_char_width * 2:
                    spaces = int(gap / avg_char_width / 1.5) 
                    line_str += " " * min(spaces, 20)
            else:
                if gap > avg_char_width * 0.8:
                    spaces = int(gap / avg_char_width)
                    line_str += " " * min(spaces, 40)
                else:
                    line_str += " "
                    
            line_str += box['text']
            last_x = box['max_x']
            
        output.append(line_str)
        
    return "\n".join(output)


def process_file_with_ocr(file_path, language, output_format, engine_name='easyocr', job_id=None, page_range=None):
    config = get_config()
    def update_progress(percentage, status):
        if job_id:
            job_progress[job_id] = {"status": status, "percentage": percentage}

    def check_controls():
        if job_id and job_id in job_controls:
            # Check for cancellation
            if job_controls[job_id].get('cancel_flag'):
                raise InterruptedError("Job cancelled by user")
            # Handle pause
            job_controls[job_id]['pause_event'].wait()

    update_progress(5, "Initializing OCR engine...")
    if engine_name == 'easyocr':
        langs = get_easyocr_langs(language)
        reader = get_reader(langs)
    elif engine_name == 'tesseract':
        lang_map = {'english': 'eng', 'bengali': 'ben', 'both': 'eng+ben'}
        tess_lang = lang_map.get(language, 'eng')
    elif engine_name == 'openai_compatible':
        openai_caps = get_openai_compat_capabilities()

    extracted_text = []
    dpi = config['OCR_DPI']

    # Handle page range
    total_pdf_pages = 0
    if is_pdf_file(file_path):
        doc = fitz.open(file_path)
        total_pdf_pages = len(doc)
        doc.close()

    page_indices = parse_page_range(page_range, total_pdf_pages) if is_pdf_file(file_path) else None
    page_images = get_document_page_images(file_path, dpi=dpi, page_indices=page_indices)
    total_pages = len(page_images)

    for idx, (page_number, img) in enumerate(page_images):
        check_controls()
        update_progress(
            5 + int((idx / total_pages) * 85),
            f"Processing page {page_number} of {total_pages}..."
        )

        try:
            if engine_name == 'easyocr':
                # EasyOCR can take PIL images directly
                import numpy as np
                img_np = np.array(img)
                result = reader.readtext(img_np, detail=1)
                text = reconstruct_layout(result)
            elif engine_name == 'tesseract':
                import pytesseract
                from pytesseract import Output
                if not configure_tesseract_command(pytesseract):
                    raise RuntimeError(
                        "Tesseract executable not found. Install Tesseract OCR "
                        "or add tesseract.exe to your PATH."
                    )
                # pytesseract can take PIL images directly
                data = pytesseract.image_to_data(img, lang=tess_lang, output_type=Output.DICT)

                ocr_results = []
                n_boxes = len(data['text'])
                for i in range(n_boxes):
                    if int(data['conf'][i]) > -1:
                        val = data['text'][i].strip()
                        if val:
                            x, y, w, h = data['left'][i], data['top'][i], data['width'][i], data['height'][i]
                            bounds = [[x, y], [x+w, y], [x+w, y+h], [x, y+h]]
                            ocr_results.append((bounds, val, data['conf'][i]))
                text = reconstruct_layout(ocr_results)
            elif engine_name == 'openai_compatible':
                text_model = openai_caps['text_model']
                vision_model = openai_caps['vision_model']

                if vision_model:
                    try:
                        text = extract_text_with_openai_compatible(img, page_number, vision_model)
                    except Exception as vision_exc:
                        if should_fallback_openai_vision(vision_exc):
                            fallback_text = extract_text_with_local_ocr_fallback(img, language)
                            if config['OPENAI_COMPAT_ENABLE_POSTPROCESS']:
                                try:
                                    text = cleanup_ocr_text_with_openai_compatible(fallback_text, text_model)
                                except Exception:
                                    text = fallback_text
                            else:
                                text = fallback_text
                        else:
                            raise
                else:
                    fallback_text = extract_text_with_local_ocr_fallback(img, language)
                    if config['OPENAI_COMPAT_ENABLE_POSTPROCESS']:
                        try:
                            text = cleanup_ocr_text_with_openai_compatible(fallback_text, text_model)
                        except Exception:
                            text = fallback_text
                    else:
                        text = fallback_text
        except Exception as e:
            text = f"[OCR Error on page {page_number}: {str(e)}]"

        extracted_text.append({
            'page': page_number,
            'text': text or ""
        })

    full_text = "\n\n".join([page['text'] for page in extracted_text])

    if output_format == 'json':
        openai_caps = get_openai_compat_capabilities()
        dataset = transform_ocr_text_to_mcq_dataset(full_text, openai_caps['text_model'])
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], f'{job_id or uuid.uuid4().hex}.json')
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(dataset, f, ensure_ascii=False, indent=2)
        preview_text = json.dumps(dataset, ensure_ascii=False, indent=2)
    elif output_format == 'txt':
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], f'{job_id or uuid.uuid4().hex}.txt')
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(full_text)
        preview_text = full_text
    elif output_format == 'docx':
        from docx import Document
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], f'{job_id or uuid.uuid4().hex}.docx')
        doc = Document()
        doc.add_heading('Extracted Text from PDF', 0)

        for page_data in extracted_text:
            doc.add_heading(f'Page {page_data["page"]}', level=1)
            for line_text in page_data['text'].split('\n'):
                p = doc.add_paragraph(line_text)
                for run in p.runs:
                    run.font.name = 'Courier New'

        doc.save(output_path)
        preview_text = full_text
    else:
        raise ValueError("Unsupported output format.")

    session['output_file'] = output_path

    return {
        'text': preview_text,
        'pages': len(extracted_text),
        'output_path': output_path
    }


@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('admin_dashboard'))

    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')

        try:
            conn = mysql.connector.connect(**app.config['DB_CONFIG'])
            cursor = conn.cursor(dictionary=True)
            cursor.execute('SELECT * FROM admins WHERE username = %s', (username,))
            user_data = cursor.fetchone()
            cursor.close()
            conn.close()

            if user_data and bcrypt.check_password_hash(user_data['password'], password):
                user = User(id=user_data['id'], username=user_data['username'],
                            name=user_data['name'], email=user_data['email'])
                login_user(user)
                return jsonify({'success': True, 'redirect': url_for('admin_dashboard')})
            else:
                return jsonify({'success': False, 'error': 'Invalid username or password'}), 401
        except Error as e:
            return jsonify({'success': False, 'error': str(e)}), 500

    return render_template('admin/login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

@app.route('/create_initial_admin')
def create_initial_admin():
    # Only allow creating an admin if the table is empty
    try:
        conn = mysql.connector.connect(**app.config['DB_CONFIG'])
        cursor = conn.cursor()
        cursor.execute('SELECT COUNT(*) FROM admins')
        count = cursor.fetchone()[0]

        if count == 0:
            username = 'admin'
            password = 'password' # The user should change this
            hashed_password = bcrypt.generate_password_hash(password).decode('utf-8')
            cursor.execute('INSERT INTO admins (username, password, name, email) VALUES (%s, %s, %s, %s)',
                         (username, hashed_password, 'Administrator', 'admin@example.com'))
            conn.commit()
            cursor.close()
            conn.close()
            return f"Initial admin created. Username: {username}, Password: {password}. Please delete this route or change the password immediately."
        else:
            cursor.close()
            conn.close()
            return "Admin already exists."
    except Error as e:
        return str(e)

@app.route('/pipelines')
def get_active_pipelines():
    try:
        conn = mysql.connector.connect(**app.config['DB_CONFIG'])
        cursor = conn.cursor(dictionary=True)
        cursor.execute('SELECT id, name FROM custom_pipelines WHERE is_active = TRUE')
        pipelines = cursor.fetchall()
        cursor.close()
        conn.close()
        return jsonify({'success': True, 'pipelines': pipelines})
    except Error as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/')
def index():
    return render_template('index.html')


@app.route('/engines')
def get_engines():
    try:
        conn = mysql.connector.connect(**app.config['DB_CONFIG'])
        cursor = conn.cursor(dictionary=True)
        cursor.execute('SELECT * FROM processing_engines WHERE is_active = TRUE ORDER BY sort_order ASC')
        engines = cursor.fetchall()
        cursor.close()
        conn.close()
        return jsonify({'success': True, 'engines': engines})
    except Error as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/upload', methods=['POST'])
def upload():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400

    file = request.files['file']

    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    # Dynamic size check
    max_mb = int(get_db_setting('MAX_UPLOAD_SIZE_MB', '1024'))
    file.seek(0, os.SEEK_END)
    file_length = file.tell()
    file.seek(0)

    if file_length > max_mb * 1024 * 1024:
        return jsonify({'error': f'File size exceeds the {max_mb}MB limit set by admin.'}), 400

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        job_id = uuid.uuid4().hex
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], f'{job_id}_{filename}')
        file.save(filepath)
        session['uploaded_file'] = filepath
        session['original_filename'] = filename
        session['job_id'] = job_id
        return jsonify({'success': True, 'filename': filename, 'job_id': job_id}), 200

    return jsonify({'error': 'Invalid file type. Upload a PDF or image file.'}), 400


@app.route('/process', methods=['POST'])
def process():
    if 'uploaded_file' not in session:
        return jsonify({'error': 'No file uploaded'}), 400

    job_id = session.get('job_id')
    if not job_id:
        return jsonify({'error': 'No job ID found'}), 400

    engine = request.form.get('engine', 'easyocr')
    pipeline_id = request.form.get('pipeline_id')
    language = request.form.get('language', 'english')
    output_format = request.form.get('format', 'txt')
    page_range = request.form.get('page_range', 'all')
    file_path = session['uploaded_file']

    # Initialize job controls
    job_controls[job_id] = {
        'pause_event': threading.Event(),
        'cancel_flag': False
    }
    job_controls[job_id]['pause_event'].set() # Initially running

    def run_process():
        try:
            job_progress[job_id] = {"status": "Starting...", "percentage": 0}

            if pipeline_id:
                result = run_custom_pipeline(file_path, pipeline_id, output_format, job_id, page_range, language)
            elif engine == 'aistudio':
                result = process_with_aistudio(file_path, output_format, job_id, page_range)
            elif engine == 'gemini':
                result = process_with_gemini(file_path, output_format, job_id, page_range)
            elif engine == 'anthropic':
                result = process_with_anthropic(file_path, output_format, job_id, page_range)
            else:
                result = process_file_with_ocr(file_path, language, output_format, engine, job_id, page_range)

            job_progress[job_id] = {
                "status": "Completed",
                "percentage": 100,
                "result": {
                    'success': True,
                    'text': result['text'],
                    'pages': result['pages'],
                    'output_file': result['output_path']
                }
            }
        except InterruptedError:
            job_progress[job_id] = {"status": "Cancelled", "percentage": 0}
        except Exception as e:
            job_progress[job_id] = {"status": "Error", "percentage": 0, "error": str(e)}
        finally:
            # Cleanup job controls after completion/error/cancel
            if job_id in job_controls:
                del job_controls[job_id]

    thread = threading.Thread(target=run_process)
    thread.start()

    return jsonify({'success': True, 'job_id': job_id}), 200


@app.route('/download/<format>')
def download(format):
    job_id = request.args.get('job_id')
    output_path = None

    if job_id and job_id in job_progress:
        job_data = job_progress[job_id]
        if job_data.get('status') == 'Completed' and 'result' in job_data:
            output_path = job_data['result'].get('output_file')

    if not output_path and 'output_file' in session:
        output_path = session['output_file']

    if not output_path or not os.path.exists(output_path):
        return jsonify({'error': 'File not found or processing not completed'}), 404

    original_name = session.get('original_filename', 'output')
    base_name = os.path.splitext(original_name)[0]

    if format == 'txt':
        return send_file(output_path, as_attachment=True, download_name=f'{base_name}.txt', mimetype='text/plain')
    elif format == 'docx':
        return send_file(output_path, as_attachment=True, download_name=f'{base_name}.docx',
                        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    elif format == 'json':
        return send_file(output_path, as_attachment=True, download_name=f'{base_name}.json', mimetype='application/json')
    elif format == 'zip':
        return send_file(output_path, as_attachment=True, download_name=f'{base_name}.zip', mimetype='application/zip')
    else:
        return jsonify({'error': 'Invalid format'}), 400


@app.route('/upload_to_db', methods=['POST'])
def upload_to_db():
    data = request.json
    if not data or 'mcqs' not in data:
        return jsonify({'error': 'Invalid JSON data'}), 400

    try:
        conn = mysql.connector.connect(**app.config['DB_CONFIG'])
        cursor = conn.cursor()
    except Error as e:
        return jsonify({'error': f'Database connection failed: {str(e)}'}), 500

    try:
        inserted_count = 0
        for item in data['mcqs']:
            # 1. Resolve Exam
            exam_name = item.get('exam') or 'Unknown Exam'
            cursor.execute('INSERT IGNORE INTO exams (name) VALUES (%s)', (exam_name,))
            cursor.execute('SELECT id FROM exams WHERE name = %s', (exam_name,))
            exam_id = cursor.fetchone()[0]

            # 2. Resolve Year
            year_val = item.get('year') or 'Unknown Year'
            cursor.execute('INSERT IGNORE INTO years (year) VALUES (%s)', (year_val,))
            cursor.execute('SELECT id FROM years WHERE year = %s', (year_val,))
            year_id = cursor.fetchone()[0]

            # 3. Insert Question
            cursor.execute('INSERT INTO questions (text) VALUES (%s)', (item['question'],))
            question_id = cursor.lastrowid

            # 4. Insert Options and collect IDs
            option_ids = []
            for opt in item['options']:
                cursor.execute('INSERT INTO options (text, type, image_base64) VALUES (%s, %s, %s)',
                             (opt.get('text'), opt.get('type', 'text'), opt.get('imageBase64')))
                option_ids.append(cursor.lastrowid)

            # 5. Insert MCQ
            answer_index = item.get('answer')
            answer_id = None
            if answer_index and 1 <= answer_index <= 4:
                answer_id = option_ids[answer_index - 1]

            cursor.execute('''
                INSERT INTO mcqs (question_id, option_ids, answer_index, answer_id, explanation, language, subject)
                VALUES (%s, %s, %s, %s, %s, %s, %s)
            ''', (
                question_id,
                json.dumps(option_ids),
                answer_index,
                answer_id,
                item.get('explanation'),
                item.get('language'),
                item.get('subject')
            ))
            mcq_id = cursor.lastrowid

            # 6. Insert into exam_questions mapping
            cursor.execute('''
                INSERT INTO exam_questions (exam_id, year_id, mcq_id)
                VALUES (%s, %s, %s)
            ''', (exam_id, year_id, mcq_id))

            inserted_count += 1

        conn.commit()
        return jsonify({'success': True, 'message': f'Successfully inserted {inserted_count} MCQs into MySQL database.'}), 200

    except Error as e:
        conn.rollback()
        return jsonify({'error': str(e)}), 500
    finally:
        if conn.is_connected():
            cursor.close()
            conn.close()


# Admin Panel Routes
@app.route('/admin')
@login_required
def admin_dashboard():
    return render_template('admin/dashboard.html')

@app.route('/admin/exams')
@login_required
def admin_exams():
    try:
        conn = mysql.connector.connect(**app.config['DB_CONFIG'])
        cursor = conn.cursor(dictionary=True)
        cursor.execute('SELECT * FROM exams ORDER BY id DESC')
        exams = cursor.fetchall()
        return render_template('admin/exams.html', exams=exams)
    except Error as e:
        return str(e), 500
    finally:
        if conn.is_connected():
            cursor.close()
            conn.close()

@app.route('/admin/years')
@login_required
def admin_years():
    try:
        conn = mysql.connector.connect(**app.config['DB_CONFIG'])
        cursor = conn.cursor(dictionary=True)
        cursor.execute('SELECT * FROM years ORDER BY id DESC')
        years = cursor.fetchall()
        return render_template('admin/years.html', years=years)
    except Error as e:
        return str(e), 500
    finally:
        if conn.is_connected():
            cursor.close()
            conn.close()

@app.route('/admin/questions')
@login_required
def admin_questions():
    try:
        conn = mysql.connector.connect(**app.config['DB_CONFIG'])
        cursor = conn.cursor(dictionary=True)
        cursor.execute('SELECT * FROM questions ORDER BY id DESC')
        questions = cursor.fetchall()
        return render_template('admin/questions.html', questions=questions)
    except Error as e:
        return str(e), 500
    finally:
        if conn.is_connected():
            cursor.close()
            conn.close()

@app.route('/admin/options')
@login_required
def admin_options():
    try:
        conn = mysql.connector.connect(**app.config['DB_CONFIG'])
        cursor = conn.cursor(dictionary=True)
        cursor.execute('SELECT * FROM options ORDER BY id DESC LIMIT 1000')
        options = cursor.fetchall()
        return render_template('admin/options.html', options=options)
    except Error as e:
        return str(e), 500
    finally:
        if conn.is_connected():
            cursor.close()
            conn.close()

def get_dir_size(path):
    total_size = 0
    if not os.path.exists(path):
        return 0
    for dirpath, dirnames, filenames in os.walk(path):
        for f in filenames:
            fp = os.path.join(dirpath, f)
            if not os.path.islink(fp):
                total_size += os.path.getsize(fp)
    return total_size

def format_size(size):
    for unit in ['B', 'KB', 'MB', 'GB']:
        if size < 1024.0:
            return f"{size:.2f} {unit}"
        size /= 1024.0
    return f"{size:.2f} TB"

@app.route('/admin/settings')
@login_required
def admin_settings():
    uploads_size = format_size(get_dir_size(app.config['UPLOAD_FOLDER']))
    outputs_size = format_size(get_dir_size(app.config['OUTPUT_FOLDER']))
    return render_template('admin/settings.html', uploads_size=uploads_size, outputs_size=outputs_size)

@app.route('/admin/add_engine', methods=['POST'])
@login_required
def add_engine():
    name = request.form.get('name')
    display_name = request.form.get('display_name')
    description = request.form.get('description')
    sort_order = request.form.get('sort_order', 0)

    if not name or not display_name:
        return jsonify({'success': False, 'error': 'Name and Display Name are required'}), 400

    try:
        conn = mysql.connector.connect(**app.config['DB_CONFIG'])
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO processing_engines (name, display_name, description, sort_order)
            VALUES (%s, %s, %s, %s)
        ''', (name, display_name, description, sort_order))
        conn.commit()
        cursor.close()
        conn.close()
        return jsonify({'success': True, 'message': 'Engine added successfully'})
    except Error as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/admin/update_setting', methods=['POST'])
@login_required
def update_setting():
    key = request.form.get('key')
    value = request.form.get('value')

    if not key:
        return jsonify({'success': False, 'error': 'Key is required'}), 400

    try:
        conn = mysql.connector.connect(**app.config['DB_CONFIG'])
        cursor = conn.cursor()
        cursor.execute('UPDATE settings SET `value` = %s WHERE `key` = %s', (value, key))
        conn.commit()
        cursor.close()
        conn.close()

        # Clear global caches if relevant settings changed
        global OpenAICompatClient, OpenAICompatCapabilities, OcrReaders
        if key.startswith('OPENAI_COMPAT_'):
            OpenAICompatClient = None
            OpenAICompatCapabilities = None
        elif key == 'EASYOCR_GPU':
            OcrReaders = {}

        return jsonify({'success': True, 'message': f'Setting {key} updated successfully'})
    except Error as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/admin/clear_folder/<folder_type>', methods=['POST'])
@login_required
def clear_folder(folder_type):
    if folder_type == 'uploads':
        folder = app.config['UPLOAD_FOLDER']
    elif folder_type == 'outputs':
        folder = app.config['OUTPUT_FOLDER']
    else:
        return jsonify({'error': 'Invalid folder type'}), 400

    try:
        count = 0
        for item in os.listdir(folder):
            item_path = os.path.join(folder, item)
            try:
                if os.path.isfile(item_path) or os.path.islink(item_path):
                    os.unlink(item_path)
                    count += 1
                elif os.path.isdir(item_path):
                    shutil.rmtree(item_path)
                    count += 1
            except Exception as e:
                print(f'Failed to delete {item_path}. Reason: {e}')

        return jsonify({'success': True, 'count': count})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/admin/clear_all_data', methods=['POST'])
@login_required
def clear_all_data():
    try:
        conn = mysql.connector.connect(**app.config['DB_CONFIG'])
        cursor = conn.cursor()

        # Disable foreign key checks to allow truncating tables in any order
        cursor.execute('SET FOREIGN_KEY_CHECKS = 0')

        tables = ['exam_questions', 'mcqs', 'questions', 'options', 'years', 'exams']
        for table in tables:
            cursor.execute(f'TRUNCATE TABLE {table}')

        cursor.execute('SET FOREIGN_KEY_CHECKS = 1')
        conn.commit()

        return jsonify({'success': True, 'message': 'All database records have been cleared.'})
    except Error as e:
        return jsonify({'error': str(e)}), 500
    finally:
        if conn.is_connected():
            cursor.close()
            conn.close()

@app.route('/admin/clear_specific_data/<table_name>', methods=['POST'])
@login_required
def clear_specific_data(table_name):
    allowed_tables = {'exams', 'years', 'questions', 'options', 'mcqs', 'exam_questions', 'ai_providers', 'custom_pipelines'}
    if table_name not in allowed_tables:
        return jsonify({'error': 'Invalid table'}), 400

    try:
        conn = mysql.connector.connect(**app.config['DB_CONFIG'])
        cursor = conn.cursor()

        # Disable foreign key checks to allow truncation if table is referenced
        cursor.execute('SET FOREIGN_KEY_CHECKS = 0')
        cursor.execute(f'TRUNCATE TABLE {table_name}')
        cursor.execute('SET FOREIGN_KEY_CHECKS = 1')
        conn.commit()

        return jsonify({'success': True, 'message': f'All records from {table_name} have been cleared.'})
    except Error as e:
        return jsonify({'error': str(e)}), 500
    finally:
        if conn.is_connected():
            cursor.close()
            conn.close()

@app.route('/admin/delete/<table_name>/<int:id>', methods=['POST'])
@login_required
def admin_delete_record(table_name, id):
    allowed_tables = {'exams', 'years', 'questions', 'options', 'mcqs', 'exam_questions', 'ai_providers', 'custom_pipelines'}
    if table_name not in allowed_tables:
        return jsonify({'error': 'Invalid table'}), 400

    try:
        conn = mysql.connector.connect(**app.config['DB_CONFIG'])
        cursor = conn.cursor()

        # Handle foreign key constraints for mcqs and exam_questions if necessary
        # For simplicity, we assume CASCADE or manual cleanup if needed.
        # Most of our tables have FKs, so deleting might fail without proper order.

        cursor.execute(f'DELETE FROM {table_name} WHERE id = %s', (id,))
        conn.commit()
        return jsonify({'success': True})
    except Error as e:
        return jsonify({'error': str(e)}), 500
    finally:
        if conn.is_connected():
            cursor.close()
            conn.close()

@app.route('/admin/bulk_delete/<table_name>', methods=['POST'])
@login_required
def admin_bulk_delete(table_name):
    allowed_tables = {'exams', 'years', 'questions', 'options', 'mcqs', 'exam_questions', 'ai_providers', 'custom_pipelines'}
    if table_name not in allowed_tables:
        return jsonify({'error': 'Invalid table'}), 400

    ids = request.json.get('ids', [])
    if not ids:
        return jsonify({'error': 'No IDs provided'}), 400

    try:
        conn = mysql.connector.connect(**app.config['DB_CONFIG'])
        cursor = conn.cursor()

        # Use safe parameter substitution for multiple IDs
        format_strings = ','.join(['%s'] * len(ids))
        cursor.execute(f'DELETE FROM {table_name} WHERE id IN ({format_strings})', tuple(ids))

        conn.commit()
        return jsonify({'success': True, 'count': cursor.rowcount})
    except Error as e:
        return jsonify({'error': str(e)}), 500
    finally:
        if conn.is_connected():
            cursor.close()
            conn.close()

@app.route('/admin/edit/<table_name>/<int:id>', methods=['POST'])
@login_required
def admin_edit_record(table_name, id):
    allowed_tables = {'exams', 'years', 'questions', 'options', 'mcqs', 'exam_questions', 'ai_providers', 'custom_pipelines'}
    if table_name not in allowed_tables:
        return jsonify({'error': 'Invalid table'}), 400

    data = request.json
    if not data:
        return jsonify({'error': 'No data provided'}), 400

    try:
        conn = mysql.connector.connect(**app.config['DB_CONFIG'])
        cursor = conn.cursor()

        # Construct dynamic UPDATE query
        fields = []
        values = []
        for key, value in data.items():
            if key != 'id':
                fields.append(f"`{key}` = %s")
                values.append(value)

        if not fields:
            return jsonify({'error': 'No fields to update'}), 400

        query = f"UPDATE `{table_name}` SET {', '.join(fields)} WHERE `id` = %s"
        values.append(id)

        cursor.execute(query, tuple(values))
        conn.commit()
        return jsonify({'success': True})
    except Error as e:
        return jsonify({'error': str(e)}), 500
    finally:
        if conn.is_connected():
            cursor.close()
            conn.close()

@app.route('/admin/mcqs')
@login_required
def admin_mcqs():
    try:
        conn = mysql.connector.connect(**app.config['DB_CONFIG'])
        cursor = conn.cursor(dictionary=True)
        cursor.execute('''
            SELECT m.*, q.text as question_text
            FROM mcqs m
            JOIN questions q ON m.question_id = q.id
            ORDER BY m.id DESC
        ''')
        mcqs = cursor.fetchall()

        # Parse option_ids JSON for each MCQ
        for mcq in mcqs:
            if mcq['option_ids']:
                try:
                    mcq['option_ids_list'] = json.loads(mcq['option_ids'])
                except:
                    mcq['option_ids_list'] = []
            else:
                mcq['option_ids_list'] = []

        return render_template('admin/mcqs.html', mcqs=mcqs)
    except Error as e:
        return str(e), 500
    finally:
        if conn.is_connected():
            cursor.close()
            conn.close()

@app.route('/admin/exam_questions')
@login_required
def admin_exam_questions():
    try:
        conn = mysql.connector.connect(**app.config['DB_CONFIG'])
        cursor = conn.cursor(dictionary=True)
        
        # 1. Fetch all mappings with basic MCQ info
        cursor.execute('''
            SELECT 
                eq.id as mapping_id,
                e.name as exam_name,
                y.year as exam_year,
                m.id as mcq_id,
                m.subject,
                m.answer_index,
                m.explanation,
                m.option_ids,
                q.text as question_text
            FROM exam_questions eq
            JOIN exams e ON eq.exam_id = e.id
            JOIN years y ON eq.year_id = y.id
            JOIN mcqs m ON eq.mcq_id = m.id
            JOIN questions q ON m.question_id = q.id
            ORDER BY eq.id DESC
        ''')
        mappings = cursor.fetchall()

        # 2. For each mapping, fetch option details
        for item in mappings:
            option_ids = []
            if item['option_ids']:
                try:
                    option_ids = json.loads(item['option_ids'])
                except:
                    pass
            
            item['options_detail'] = []
            if option_ids:
                format_strings = ','.join(['%s'] * len(option_ids))
                cursor.execute(f"SELECT * FROM options WHERE id IN ({format_strings})", tuple(option_ids))
                options = cursor.fetchall()
                
                # Sort options to match the order in option_ids
                opt_map = {opt['id']: opt for opt in options}
                item['options_detail'] = [opt_map[oid] for oid in option_ids if oid in opt_map]

        # 3. Fetch exams and years for the edit modal
        cursor.execute("SELECT * FROM exams ORDER BY name ASC")
        exams = cursor.fetchall()
        cursor.execute("SELECT * FROM years ORDER BY year DESC")
        years = cursor.fetchall()

        cursor.close()
        conn.close()
        
        return render_template('admin/exam_questions.html', 
                             data=mappings, 
                             exams=exams, 
                             years=years)
    except Error as e:
        return str(e), 500
    finally:
        if 'conn' in locals() and conn.is_connected():
            conn.close()


@app.route('/admin/pipelines')
@login_required
def admin_pipelines():
    try:
        conn = mysql.connector.connect(**app.config['DB_CONFIG'])
        cursor = conn.cursor(dictionary=True)

        # Fetch settings (Moved from settings route)
        cursor.execute('SELECT * FROM settings ORDER BY category, `key`')
        db_settings = cursor.fetchall()

        # Fetch engines (Moved from settings route)
        cursor.execute('SELECT * FROM processing_engines ORDER BY sort_order ASC')
        db_engines = cursor.fetchall()

        # Fetch providers
        cursor.execute('SELECT * FROM ai_providers ORDER BY name ASC')
        providers = cursor.fetchall()

        # Fetch pipelines
        cursor.execute('''
            SELECT cp.*,
                   op.name as ocr_provider_name,
                   sp.name as structure_provider_name
            FROM custom_pipelines cp
            LEFT JOIN ai_providers op ON cp.ocr_provider_id = op.id
            LEFT JOIN ai_providers sp ON cp.structure_provider_id = sp.id
            ORDER BY cp.name ASC
        ''')
        pipelines = cursor.fetchall()

        cursor.close()
        conn.close()
        return render_template('admin/pipelines.html',
                             providers=providers,
                             pipelines=pipelines,
                             settings=db_settings,
                             engines=db_engines)
    except Error as e:
        return str(e), 500

@app.route('/admin/add_provider', methods=['POST'])
@login_required
def add_provider():
    name = request.form.get('name')
    p_type = request.form.get('type')
    base_url = request.form.get('base_url')
    api_key = request.form.get('api_key')

    try:
        conn = mysql.connector.connect(**app.config['DB_CONFIG'])
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO ai_providers (name, type, base_url, api_key)
            VALUES (%s, %s, %s, %s)
        ''', (name, p_type, base_url, api_key))
        conn.commit()
        cursor.close()
        conn.close()
        return jsonify({'success': True})
    except Error as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/admin/add_custom_pipeline', methods=['POST'])
@login_required
def add_custom_pipeline():
    name = request.form.get('name')
    ocr_provider_id = request.form.get('ocr_provider_id')
    ocr_model = request.form.get('ocr_model')
    structure_provider_id = request.form.get('structure_provider_id')
    structure_model = request.form.get('structure_model')

    try:
        conn = mysql.connector.connect(**app.config['DB_CONFIG'])
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO custom_pipelines (name, ocr_provider_id, ocr_model, structure_provider_id, structure_model)
            VALUES (%s, %s, %s, %s, %s)
        ''', (name, ocr_provider_id, ocr_model, structure_provider_id, structure_model))
        conn.commit()
        cursor.close()
        conn.close()
        return jsonify({'success': True})
    except Error as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/admin/toggle_pipeline/<int:id>', methods=['POST'])
@login_required
def toggle_pipeline(id):
    try:
        conn = mysql.connector.connect(**app.config['DB_CONFIG'])
        cursor = conn.cursor()
        cursor.execute('UPDATE custom_pipelines SET is_active = NOT is_active WHERE id = %s', (id,))
        conn.commit()
        cursor.close()
        conn.close()
        return jsonify({'success': True})
    except Error as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/admin/profile')
@login_required
def admin_profile():
    return render_template('admin/profile.html')

@app.route('/admin/change_password', methods=['POST'])
@login_required
def admin_change_password():
    current_password = request.form.get('current_password')
    new_password = request.form.get('new_password')
    confirm_password = request.form.get('confirm_password')

    if not current_password or not new_password or not confirm_password:
        return jsonify({'success': False, 'error': 'Missing fields'}), 400

    if new_password != confirm_password:
        return jsonify({'success': False, 'error': 'New passwords do not match'}), 400

    try:
        conn = mysql.connector.connect(**app.config['DB_CONFIG'])
        cursor = conn.cursor(dictionary=True)
        cursor.execute('SELECT password FROM admins WHERE id = %s', (current_user.id,))
        user_data = cursor.fetchone()

        if not user_data or not bcrypt.check_password_hash(user_data['password'], current_password):
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'error': 'Incorrect current password'}), 401

        hashed_password = bcrypt.generate_password_hash(new_password).decode('utf-8')
        cursor.execute('UPDATE admins SET password = %s WHERE id = %s', (hashed_password, current_user.id))
        conn.commit()
        cursor.close()
        conn.close()

        return jsonify({'success': True, 'message': 'Password updated successfully'})
    except Error as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/cleanup', methods=['POST'])
def cleanup():
    if 'uploaded_file' in session:
        try:
            if os.path.exists(session['uploaded_file']):
                os.remove(session['uploaded_file'])
        except:
            pass
        session.pop('uploaded_file', None)

    if 'output_file' in session:
        try:
            if os.path.exists(session['output_file']):
                os.remove(session['output_file'])
        except:
            pass
        session.pop('output_file', None)

    return jsonify({'success': True}), 200


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
